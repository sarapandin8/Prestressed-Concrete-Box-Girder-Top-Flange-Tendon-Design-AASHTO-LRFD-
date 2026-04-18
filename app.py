"""
PSC Box Girder — Top Flange Transverse Design  (v4 — Navy UI Edition)
AASHTO LRFD Bridge Design Specifications  |  1.0 m transverse strip

Changes from v3:
  [UI-1] Navy sidebar with high-contrast white/light text
  [UI-2] All input labels, captions, expander headers styled for readability
  [UI-3] Prominent "Export Report (.docx)" button in both sidebar and main area
  [UI-4] Summary metric cards in main area header
  [UI-5] Result tabs with improved header styling

Calculation logic: UNCHANGED from v3 fixed.
"""

import math, datetime, json, os
from io import BytesIO

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ─────────────────────────────────────────────────────────────────────────────
# 1.  CONFIG & SESSION STATE INITIALIZATION
# ─────────────────────────────────────────────────────────────────────────────

# โหลดรูป box girder มาเป็นไอคอน
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
try:
    box_girder_icon = Image.open(os.path.join(_BASE_DIR, "image_7a2891.png"))
except FileNotFoundError:
    box_girder_icon = "🏗️"

st.set_page_config(
    layout="wide",
    page_title="PSC Box Girder — Top Flange Design",
    page_icon=box_girder_icon,
)
# ─────────────────────────────────────────────────────────────────────────────
# CUSTOM CSS — Navy Sidebar + High-Contrast UI
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Google Font ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* ═══════════════════════════════════════════════
   SIDEBAR — Deep Navy Theme
═══════════════════════════════════════════════ */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0a1628 0%, #0d1f3c 40%, #112244 100%) !important;
    border-right: 1px solid #1e3a5f !important;
}
[data-testid="stSidebar"] > div:first-child {
    background: transparent !important;
}

/* Sidebar text — bright white for max contrast */
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] .stMarkdown p,
[data-testid="stSidebar"] div[data-testid="stMarkdownContainer"] p {
    color: #e8f4fd !important;
    font-weight: 500 !important;
}

/* Sidebar headers */
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
    font-weight: 700 !important;
}

/* Sidebar input fields */
[data-testid="stSidebar"] input[type="number"],
[data-testid="stSidebar"] input[type="text"],
[data-testid="stSidebar"] select,
[data-testid="stSidebar"] textarea {
    background-color: #1a3050 !important;
    color: #ffffff !important;
    border: 1px solid #2d5a8e !important;
    border-radius: 6px !important;
    font-weight: 500 !important;
}
[data-testid="stSidebar"] input[type="number"]:focus,
[data-testid="stSidebar"] input[type="text"]:focus {
    border-color: #4da6ff !important;
    box-shadow: 0 0 0 2px rgba(77,166,255,0.25) !important;
}

/* Sidebar selectbox */
[data-testid="stSidebar"] [data-testid="stSelectbox"] > div > div {
    background-color: #1a3050 !important;
    color: #ffffff !important;
    border: 1px solid #2d5a8e !important;
}

/* Sidebar number input label */
[data-testid="stSidebar"] .stNumberInput label,
[data-testid="stSidebar"] .stTextInput label,
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stSlider label,
[data-testid="stSidebar"] .stFileUploader label {
    color: #93c5fd !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    text-transform: uppercase !important;
}

/* Sidebar number input arrows */
[data-testid="stSidebar"] input[type="number"]::-webkit-inner-spin-button {
    filter: invert(1) !important;
}

/* Sidebar expander header */
[data-testid="stSidebar"] [data-testid="stExpander"] summary {
    background-color: #162d50 !important;
    border: 1px solid #1e3f6e !important;
    border-radius: 8px !important;
    color: #ffffff !important;
    font-weight: 600 !important;
    padding: 8px 12px !important;
}
[data-testid="stSidebar"] [data-testid="stExpander"] summary:hover {
    background-color: #1e3f6e !important;
}
[data-testid="stSidebar"] [data-testid="stExpander"] summary p,
[data-testid="stSidebar"] [data-testid="stExpander"] summary span {
    color: #ffffff !important;
    font-weight: 600 !important;
}

/* Expander content area */
[data-testid="stSidebar"] [data-testid="stExpander"] > div:last-child {
    background-color: rgba(10, 30, 60, 0.4) !important;
    border: 1px solid #1e3a5f !important;
    border-top: none !important;
    border-radius: 0 0 8px 8px !important;
    padding: 8px !important;
}

/* Sidebar caption */
[data-testid="stSidebar"] [data-testid="stCaptionContainer"] p {
    color: #7eb3e6 !important;
    font-size: 0.73rem !important;
    font-weight: 400 !important;
    text-transform: none !important;
}

/* Sidebar slider */
[data-testid="stSidebar"] .stSlider [role="slider"] {
    background-color: #4da6ff !important;
    border: 2px solid #ffffff !important;
}
[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] > div:first-child {
    background-color: #2d5a8e !important;
}

/* ── Sidebar file uploader — full navy theme ── */

/* Label above the dropzone */
[data-testid="stSidebar"] [data-testid="stFileUploaderLabel"],
[data-testid="stSidebar"] .stFileUploader label,
[data-testid="stSidebar"] .stFileUploader > label {
    color: #93c5fd !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    text-transform: uppercase !important;
}
/* Dropzone container - แก้ Gradient เขียว + ทุบ section ขาว */
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] {
    background: linear-gradient(135deg, #1e6030 0%, #2d8a4a 100%)!important;
    border: 2px dashed #2d5a8e!important;
    border-radius: 8px!important;
    padding: 12px!important;
}
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"]:hover {
    border-color: #4da6ff!important;
    background: linear-gradient(135deg, #2d8a4a 0%, #3ba55d 100%)!important;
}

/* ตัวการ: section ข้างในที่พื้นขาว จับทำให้โปร่งใส */
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] section.st-emotion-cache-nd5v17,
[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] {
    background: transparent!important;
    background-color: transparent!important;
}

/* ตัวหนังสือทั้งหมดเป็นสีขาว */
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] *,
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] span,
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] p,
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] small,
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] div {
    color: #ffffff!important;
    font-weight: 500!important;
}

/* ปุ่ม Upload - ทำให้เขียวทึบเหมือน Save Project */
[data-testid="stSidebar"] [data-testid="stBaseButton-secondary"],
[data-testid="stSidebar"] button.st-emotion-cache-1dn7sub {
    background: linear-gradient(135deg, #1e6030 0%, #2d8a4a 100%)!important;
    background-color: #1e6030!important;
    background-image: linear-gradient(135deg, #1e6030 0%, #2d8a4a 100%)!important;
    color: #ffffff!important;
    border: 1px solid #2d8a4a!important;
    border-radius: 6px!important;
    box-shadow: none!important;
    font-weight: 600!important;
}

/* เคลียร์ทุกอย่างข้างในปุ่มให้โปร่งใส */
[data-testid="stSidebar"] [data-testid="stBaseButton-secondary"] *,
[data-testid="stSidebar"] button.st-emotion-cache-1dn7sub * {
    background: transparent!important;
    background-color: transparent!important;
    background-image: none!important;
    color: #ffffff!important;
}

/* ตอน hover ให้สว่างขึ้นเหมือน Save Project */
[data-testid="stSidebar"] [data-testid="stBaseButton-secondary"]:hover,
[data-testid="stSidebar"] button.st-emotion-cache-1dn7sub:hover {
    background: linear-gradient(135deg, #2d8a4a 0%, #3ba55d 100%)!important;
    background-color: #2d8a4a!important;
    border-color: #4da6ff!important;
}

/* "Browse files" button inside dropzone */
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] button {
    background-color: #1e3f6e !important;
    color: #ffffff !important;
    border: 1px solid #2d5a8e !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.80rem !important;
    padding: 5px 14px !important;
    transition: all 0.2s ease !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] button:hover {
    background-color: #2d5a8e !important;
    border-color: #4da6ff !important;
    color: #ffffff !important;
}

/* Helper text "200MB per file • JSON" */
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] small,
[data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] small {
    color: #5b8db8 !important;
    font-size: 0.72rem !important;
    font-weight: 400 !important;
}

/* Uploaded file chip / name shown after upload */
[data-testid="stSidebar"] [data-testid="stFileUploaderFile"] {
    background-color: #162d50 !important;
    border: 1px solid #2d5a8e !important;
    border-radius: 6px !important;
    color: #e8f4fd !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderFile"] span,
[data-testid="stSidebar"] [data-testid="stFileUploaderFileName"] {
    color: #e8f4fd !important;
    font-weight: 500 !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderFile"] button {
    background: transparent !important;
    color: #93c5fd !important;
    border: none !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderFile"] button:hover {
    color: #ff6b6b !important;
}

/* Sidebar horizontal rule */
[data-testid="stSidebar"] hr {
    border-color: #1e3a5f !important;
    border-width: 1px !important;
    margin: 10px 0 !important;
}

/* Sidebar info/success/error boxes */
[data-testid="stSidebar"] [data-testid="stAlert"] {
    border-radius: 6px !important;
}
[data-testid="stSidebar"] .stSuccess {
    background-color: #0d3320 !important;
    border-color: #1a6b3a !important;
    color: #4ade80 !important;
}
[data-testid="stSidebar"] .stError {
    background-color: #3b0d0d !important;
    border-color: #7f1d1d !important;
    color: #fca5a5 !important;
}

/* ── SAVE button (Download) — Gold accent in sidebar ── */
[data-testid="stSidebar"] [data-testid="stDownloadButton"] button {
    background: linear-gradient(135deg, #1e6030 0%, #2d8a4a 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
    font-size: 0.85rem !important;
    padding: 10px 16px !important;
    letter-spacing: 0.03em !important;
    transition: all 0.2s ease !important;
}
[data-testid="stSidebar"] [data-testid="stDownloadButton"] button:hover {
    background: linear-gradient(135deg, #2d8a4a 0%, #38a85d 100%) !important;
    box-shadow: 0 4px 15px rgba(45, 138, 74, 0.4) !important;
    transform: translateY(-1px) !important;
}

/* ═══════════════════════════════════════════════
   MAIN AREA
═══════════════════════════════════════════════ */

/* Main app background */
[data-testid="stAppViewContainer"] > .main {
    background-color: #f0f4f8 !important;
}

/* Subheaders in main */
.stApp h2, .stApp h3 {
    color: #0d1f3c !important;
    font-weight: 700 !important;
}

/* Metric containers */
[data-testid="metric-container"] {
    background: linear-gradient(135deg, #0d1f3c, #1a3a6e) !important;
    border-radius: 10px !important;
    padding: 14px 18px !important;
    border: 1px solid #2d5a8e !important;
}
[data-testid="metric-container"] label {
    color: #93c5fd !important;
    font-size: 0.75rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #ffffff !important;
    font-size: 1.6rem !important;
    font-weight: 800 !important;
}
[data-testid="metric-container"] [data-testid="stMetricDelta"] {
    color: #4ade80 !important;
    font-weight: 600 !important;
}

/* Tab styling */
[data-testid="stTabs"] [role="tab"] {
    background-color: #e8f0f8 !important;
    color: #0d1f3c !important;
    border-radius: 8px 8px 0 0 !important;
    font-weight: 600 !important;
    border: 1px solid #c5d5e8 !important;
    border-bottom: none !important;
    padding: 8px 16px !important;
    font-size: 0.83rem !important;
}
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {
    background: linear-gradient(180deg, #0d1f3c, #1a3a6e) !important;
    color: #ffffff !important;
    border-color: #1a3a6e !important;
}
[data-testid="stTabs"] [role="tabpanel"] {
    background-color: #ffffff !important;
    border-radius: 0 0 10px 10px !important;
    border: 1px solid #c5d5e8 !important;
    padding: 20px !important;
}

/* Data editor */
[data-testid="stDataEditor"] {
    border-radius: 8px !important;
    overflow: hidden !important;
    border: 1px solid #c5d5e8 !important;
}

/* Info / warning / success / error in main */
[data-testid="stAlert"][data-baseweb="notification"] {
    border-radius: 8px !important;
    font-weight: 500 !important;
}

/* ── EXPORT REPORT button — Prominent amber/gold ── */
.export-btn-container button {
    background: linear-gradient(135deg, #b45309 0%, #d97706 50%, #f59e0b 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 800 !important;
    font-size: 1.0rem !important;
    padding: 14px 28px !important;
    letter-spacing: 0.04em !important;
    box-shadow: 0 4px 20px rgba(217, 119, 6, 0.45) !important;
    transition: all 0.25s ease !important;
}
.export-btn-container button:hover {
    background: linear-gradient(135deg, #d97706 0%, #f59e0b 50%, #fbbf24 100%) !important;
    box-shadow: 0 6px 28px rgba(251, 191, 36, 0.55) !important;
    transform: translateY(-2px) !important;
}

/* ── Column headers in dataframes ── */
.stDataFrame thead th {
    background-color: #0d1f3c !important;
    color: #ffffff !important;
    font-weight: 700 !important;
}

/* Subheader in main content */
.block-container h2 {
    padding-top: 0.5rem !important;
    border-bottom: 2px solid #c5d5e8 !important;
    padding-bottom: 6px !important;
}

</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
def _make_float_df(data: dict) -> pd.DataFrame:
    """สร้าง DataFrame และบังคับทุกคอลัมน์เป็น float64 กันเคสเพิ่มแถวแล้วกลายเป็น int"""
    df = pd.DataFrame({col: pd.Series(vals, dtype="float64") for col, vals in data.items()})
    return df
DEFAULT_SCALARS = dict(
    width=12.0, cl_lweb=2.0, cl_rweb=10.0,
    fc=45.0, fci=36.0, fpu=1860.0, fpy_ratio=0.90,
    aps_strand=140.0, duct_dia_mm=70.0,
    num_tendon=1, n_strands=5,
    fpi_ratio=0.75,
    t0=3, RH=75,
    anch_slip_mm=6.0,
    phi_flex=1.00, phi_shear=0.90,
    proj_name="Box Girder Design", doc_no="CALC-STR-001",
    eng_name="Engineer Name", chk_name="Checker Name",
)
DEFAULT_TABLES = dict(
    df_thickness={"x (m)": [0,1,2,3,4,5,6,7,8,9,10,11,12],
                  "t (m)":  [0.25]*13},
    df_tendon=   {"x (m)": [0,1,2,3,4,5,6,7,8,9,10,11,12],
                  "z_top (m)": [0.10]*13},
    df_load={
        "x (m)":         [0,1,2,3,4,5,6,7,8,9,10,11,12],
        "M_DL (kNm/m)":  [0]*13,  "V_DL (kN/m)":  [0]*13,
        "M_SDL (kNm/m)": [0]*13,  "V_SDL (kN/m)": [0]*13,
        "M_LL (kNm/m)":  [0]*13,  "V_LL (kN/m)":  [0]*13,
    },
)

# วางต่อตรงนี้เลย - จำค่า default ไว้ตอนเปิดแอป
if "_default_thk" not in st.session_state:
    st.session_state["_default_thk"] = st.session_state["thk_src"].copy()
if "_default_tdn" not in st.session_state:
    st.session_state["_default_tdn"] = st.session_state["tdn_src"].copy()
if "_default_ld" not in st.session_state:
    st.session_state["_default_ld"] = st.session_state["ld_src"].copy()

for k, v in DEFAULT_SCALARS.items():
    if k not in st.session_state:
        st.session_state[k] = v

_TABLE_SRC = {"thk_src": "df_thickness", "tdn_src": "df_tendon", "ld_src": "df_load"}
for src_key, tbl_key in _TABLE_SRC.items():
    if src_key not in st.session_state:
        st.session_state[src_key] = _make_float_df(DEFAULT_TABLES[tbl_key])

if "_tbl_ver" not in st.session_state:
    st.session_state["_tbl_ver"] = 0
if "_loaded_hash" not in st.session_state:
    st.session_state["_loaded_hash"] = None


# ─────────────────────────────────────────────────────────────────────────────
# 2.  SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    # แสดงรูป Box Girder ในหน้าแอป
    if isinstance(box_girder_icon, Image.Image):
        st.image(box_girder_icon, use_container_width=True)

    # Logo / Header
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #0a1628, #162d50);
        border-radius: 10px;
        padding: 14px 16px;
        margin-bottom: 12px;
        border: 1px solid #1e3f6e;
    ">
        <div style="font-size:1.3rem; font-weight:800; color:#ffffff; line-height:1.2;">
            🏗️ PSC Box Girder
        </div>
        <div style="font-size:0.68rem; color:#93c5fd; font-weight:600;
                    letter-spacing:0.12em; margin-top:4px; text-transform:uppercase;">
            AASHTO LRFD  ·  1.0 m Transverse Strip
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── 💾 SAVE / 📂 OPEN ────────────────────────────────────────────────────
    with st.expander("💾  Save  /  📂  Open Project", expanded=True):
        def _tbl_to_dict(cur_key, src_key):
            df = st.session_state.get(cur_key,
                 st.session_state.get(src_key, pd.DataFrame()))
            if not isinstance(df, pd.DataFrame):
                try:    df = pd.DataFrame(df)
                except: return {}
            for col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            df = df.dropna(how="all")
            return df.to_dict(orient="list") if not df.empty else {}

        _save_data = {
            "scalars": {k: st.session_state[k] for k in DEFAULT_SCALARS},
            "tables": {
                "df_thickness": _tbl_to_dict("_cur_thk", "thk_src"),
                "df_tendon":    _tbl_to_dict("_cur_tdn", "tdn_src"),
                "df_load":      _tbl_to_dict("_cur_ld",  "ld_src"),
            },
        }
        _json_bytes = json.dumps(_save_data, indent=2, ensure_ascii=False).encode("utf-8")
        _fname = f"{st.session_state.proj_name.replace(' ','_')}_{st.session_state.doc_no}.json"
        st.download_button(
            label="💾  Save Project  (.json)",
            data=_json_bytes, file_name=_fname,
            mime="application/json", use_container_width=True,
        )
        st.caption("ตั้ง Chrome: Settings → Downloads → 'Ask where to save'")
        st.markdown("---")

        uploaded_file = st.file_uploader(
            "📂  Open Project  (.json)", type="json",
            key="proj_uploader",
            help="เลือกไฟล์ .json ที่เคย Save ไว้",
        )
        if uploaded_file is not None:
            _raw   = uploaded_file.getvalue()
            _fhash = hash(_raw)
            if st.session_state["_loaded_hash"] != _fhash:
                try:
                    loaded = json.loads(_raw.decode("utf-8"))
                    for k, v in loaded.get("scalars", {}).items():
                        if k in DEFAULT_SCALARS:
                            dv = DEFAULT_SCALARS[k]
                            st.session_state[k] = (
                                int(v)   if isinstance(dv, int)   else
                                float(v) if isinstance(dv, float) else str(v)
                            )
                    _lmap = {"df_thickness":"thk_src","df_tendon":"tdn_src","df_load":"ld_src"}
                    for tbl_key, src_key in _lmap.items():
                        raw_tbl = loaded.get("tables", {}).get(tbl_key)
                        if raw_tbl:
                            ndf = pd.DataFrame(raw_tbl)
                            for col in ndf.columns:
                                ndf[col] = pd.to_numeric(ndf[col], errors="coerce")
                            st.session_state[src_key] = ndf.dropna(how="all")
                    st.session_state["_tbl_ver"] += 1
                    for k in ["_cur_thk", "_cur_tdn", "_cur_ld"]:
                        st.session_state.pop(k, None)
                    st.session_state["_loaded_hash"] = _fhash
                    st.success("✅  Project loaded successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌  Load error: {e}")

    # ── Materials & Section ──────────────────────────────────────────────────
    with st.expander("📐  Materials & Section", expanded=True):
        width       = st.number_input("Total Flange Width (m)",   min_value=1.0, key="width")
        fc          = st.number_input("f'c  Service (MPa)",       min_value=20.0, key="fc")
        fci         = st.number_input("f'ci Transfer (MPa)",      min_value=15.0, key="fci")
        fpu         = st.number_input("fpu (MPa)",                key="fpu")
        fpy_opts = [0.90, 0.85]
        if st.session_state.fpy_ratio not in fpy_opts:
            st.session_state.fpy_ratio = 0.90
        fpy_ratio   = st.selectbox("fpy/fpu", fpy_opts, key="fpy_ratio",
                                   help="Low-relaxation=0.90  |  Stress-relieved=0.85")
        aps_strand  = st.number_input("Aps per strand (mm²)",     key="aps_strand")
        duct_dia_mm = st.number_input("Duct diameter (mm)",       min_value=20.0, key="duct_dia_mm")

    # ── Web Geometry ─────────────────────────────────────────────────────────
    with st.expander("🟦  Web Geometry", expanded=True):
        st.caption("ระบุตำแหน่ง Centerline ของ Web ซ้าย-ขวา จากขอบซ้ายของ Flange")
        col_wl, col_wr = st.columns(2)
        cl_lweb = col_wl.number_input("CL. L.Web (m)", min_value=0.0, step=0.05, key="cl_lweb")
        cl_rweb = col_wr.number_input("CL. R.Web (m)", min_value=0.0, step=0.05, key="cl_rweb")
        st.info(f"Span = **{(cl_rweb-cl_lweb)*1000:.0f} mm**  |  "
                f"L.Web = {cl_lweb*1000:.0f} mm  |  R.Web = {cl_rweb*1000:.0f} mm")

    # ── Prestressing Force ───────────────────────────────────────────────────
    with st.expander("🔩  Prestressing Force", expanded=True):
        num_tendon = st.number_input("Tendons per 1 m strip", min_value=1, key="num_tendon")
        n_strands  = st.number_input("Strands per tendon",    min_value=1, key="n_strands")
        fpi_ratio  = st.slider("Jacking stress  fpi/fpu", 0.70, 0.80, key="fpi_ratio",
                               help="Standard = 0.75 fpu  (AASHTO 5.9.2.2)")

    # ── Prestress Losses ─────────────────────────────────────────────────────
    with st.expander("📉  Prestress Loss Parameters", expanded=True):
        st.caption("คำนวณ Loss ตาม AASHTO LRFD 5.9.3 อัตโนมัติ")
        t0_val   = st.number_input("Age at Transfer  t₀ (days)", min_value=1, key="t0")
        rh_val   = st.number_input("Relative Humidity  RH (%)", min_value=30, max_value=100, key="RH")
        anch_val = st.number_input("Anchorage Slip  Δ (mm)", value=6.0, min_value=0.0,
                                   key="anch_slip_mm")
        st.caption("Constants: μ=0.20, K=0.0066 rad/m, Ep=197,000 MPa")

    # ── Resistance Factors ───────────────────────────────────────────────────
    with st.expander("⚖️  Resistance Factors φ"):
        phi_flex  = st.number_input("φ  Flexure", min_value=0.75, max_value=1.00, key="phi_flex")
        phi_shear = st.number_input("φ  Shear",   min_value=0.70, max_value=1.00, key="phi_shear")

    # ── Report Info ──────────────────────────────────────────────────────────
    st.markdown("""
    <div style="color:#93c5fd; font-size:0.72rem; font-weight:700;
                letter-spacing:0.1em; text-transform:uppercase;
                padding:10px 0 4px 0;">
        📄  Report Information
    </div>
    """, unsafe_allow_html=True)
    proj_name = st.text_input("Project Name", key="proj_name")
    doc_no    = st.text_input("Document No.", key="doc_no")
    eng_name  = st.text_input("Prepared by",  key="eng_name")
    chk_name  = st.text_input("Checked by",   key="chk_name")


# ─────────────────────────────────────────────────────────────────────────────
# 3.  MAIN HEADER
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="
    background: linear-gradient(135deg, #0a1628 0%, #0d2044 50%, #112244 100%);
    border-radius: 12px;
    padding: 22px 28px 18px;
    margin-bottom: 20px;
    border: 1px solid #1e3a5f;
    box-shadow: 0 4px 20px rgba(10,22,40,0.4);
">
    <div style="font-size:1.5rem; font-weight:800; color:#ffffff; line-height:1.2;">
        🏗️  PSC Box Girder — Top Flange Transverse Design
    </div>
    <div style="font-size:0.82rem; color:#93c5fd; margin-top:6px; font-weight:500;">
        AASHTO LRFD Bridge Design Specifications  ·  1.0 m transverse strip  ·
        Compression (−)  Tension (+)  ·  +M = sagging
    </div>
</div>
""", unsafe_allow_html=True)
# ─────────────────────────────────────────────────────────────────────────────
# 3. DATA EDITORS
# ─────────────────────────────────────────────────────────────────────────────
_v = st.session_state["_tbl_ver"]

# CSS ปุ่ม Primary สีกรมท่า + ปุ่ม Secondary สีเทา
st.markdown("""
<style>
div[data-testid="stButton"] > button[kind="primary"] {
    background-color: #0d1f3c!important;
    border: 1px solid #0d1f3c!important;
    color: #ffffff!important;
    font-weight: 600!important;
    border-radius: 8px!important;
    transition: all 0.2s ease!important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    background-color: #1a2f4d!important;
    border: 1px solid #1a2f4d!important;
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(13, 31, 60, 0.3)!important;
}
div[data-testid="stButton"] > button[kind="primary"]:active {
    background-color: #081629!important;
    border: 1px solid #081629!important;
}
div[data-testid="stButton"] > button[kind="secondary"] {
    border-radius: 8px!important;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style="
    background: #ffffff;
    border-radius: 10px;
    padding: 16px 20px;
    margin-bottom: 16px;
    border: 1px solid #c5d5e8;
">
    <div style="font-size:1.0rem; font-weight:700; color:#0d1f3c; margin-bottom:4px;">
        📊 Input Station Tables
    </div>
    <div style="font-size:0.8rem; color:#4a6080;">
        แก้ข้อมูลในตารางได้หลายช่อง แล้วกดปุ่ม "Update Data" เพื่ออัปเดท | Hover ที่หัวตารางเพื่อดูคำอธิบาย
    </div>
</div>
""", unsafe_allow_html=True)

# แถวบน: Thickness + Tendon
c1, c2 = st.columns(2)
with c1:
    st.subheader("📏 Flange Thickness t(x)")
    df_thk_edit = st.data_editor(
        st.session_state["thk_src"].astype("float64"),
        num_rows="dynamic",
        key="_tmp_thk",
        use_container_width=True,
        column_config={
            "x (m)": st.column_config.NumberColumn(
                "x (m)",
                help="ระยะจากกึ่งกลางสะพาน x=0 ถึงขอบ B/2",
                format="%.3f"
            ),
            "t (m)": st.column_config.NumberColumn(
                "t (m)",
                help="ความหนา Top flange ที่ตำแหน่ง x นั้นๆ หน่วยเมตร",
                format="%.3f"
            ),
        }
    )

    b1, b2 = st.columns(2)
    with b1:
        if st.button("🔄 Update Thickness", key="btn_update_thk", use_container_width=True, type="primary"):
            st.session_state["thk_src"] = df_thk_edit.astype("float64")
            st.success("✅ Thickness data updated")
            st.rerun()
    with b2:
        if st.button("↩ Reset", key="btn_reset_thk", use_container_width=True):
            st.session_state["thk_src"] = st.session_state["_default_thk"].copy()
            st.success("↩ Reset to default")
            st.rerun()

with c2:
    st.subheader("🔩 Tendon Profile z(x) [from top face]")
    df_tdn_edit = st.data_editor(
        st.session_state["tdn_src"].astype("float64"),
        num_rows="dynamic",
        key="_tmp_tdn",
        use_container_width=True,
        column_config={
            "x (m)": st.column_config.NumberColumn(
                "x (m)",
                help="ระยะจากกึ่งกลางสะพาน x=0 ถึงขอบ B/2",
                format="%.3f"
            ),
            "z_top (m)": st.column_config.NumberColumn(
                "z_top (m)",
                help="ระยะจากผิวบน Flange ถึง CG ของลวดอัดแรง +ลงล่าง หน่วยเมตร",
                format="%.3f"
            ),
        }
    )

    b3, b4 = st.columns(2)
    with b3:
        if st.button("🔄 Update Tendon", key="btn_update_tdn", use_container_width=True, type="primary"):
            st.session_state["tdn_src"] = df_tdn_edit.astype("float64")
            st.success("✅ Tendon data updated")
            st.rerun()
    with b4:
        if st.button("↩ Reset", key="btn_reset_tdn", use_container_width=True):
            st.session_state["tdn_src"] = st.session_state["_default_tdn"].copy()
            st.success("↩ Reset to default")
            st.rerun()

st.divider()

# แถวล่าง: Loads กว้างเต็มจอ
st.subheader("📦 Loads per 1 m strip")
df_ld_edit = st.data_editor(
    st.session_state["ld_src"].astype("float64"),
    num_rows="dynamic",
    key="_tmp_ld",
    use_container_width=True,
    height=350,
    column_config={
        "x (m)": st.column_config.NumberColumn("x (m)", help="ตำแหน่งตามแนวขวาง", format="%.3f"),
        "M_DL (kNm/m)": st.column_config.NumberColumn("M_DL (kNm/m)", help="โมเมนต์จาก Dead Load คอนกรีต", format="%.2f"),
        "V_DL (kN/m)": st.column_config.NumberColumn("V_DL (kN/m)", help="แรงเฉือนจาก Dead Load คอนกรีต", format="%.2f"),
        "M_SDL (kNm/m)": st.column_config.NumberColumn("M_SDL (kNm/m)", help="โมเมนต์จาก Superimposed Dead Load", format="%.2f"),
        "V_SDL (kN/m)": st.column_config.NumberColumn("V_SDL (kN/m)", help="แรงเฉือนจาก Superimposed Dead Load", format="%.2f"),
        "M_LL (kNm/m)": st.column_config.NumberColumn("M_LL (kNm/m)", help="โมเมนต์จาก Live Load", format="%.2f"),
        "V_LL (kN/m)": st.column_config.NumberColumn("V_LL (kN/m)", help="แรงเฉือนจาก Live Load", format="%.2f"),
    }
)

b5, b6 = st.columns([3, 1])
with b5:
    if st.button("🔄 Update Load Data", key="btn_update_ld", use_container_width=True, type="primary"):
        st.session_state["ld_src"] = df_ld_edit.astype("float64")
        st.success("✅ Load data updated")
        st.rerun()
with b6:
    if st.button("↩ Reset", key="btn_reset_ld", use_container_width=True):
        st.session_state["ld_src"] = st.session_state["_default_ld"].copy()
        st.success("↩ Reset to default")
        st.rerun()

# ตัวแปรที่ใช้ Save/Load ให้อ่านจาก source เสมอ
st.session_state["_cur_thk"] = st.session_state["thk_src"]
st.session_state["_cur_tdn"] = st.session_state["tdn_src"]
st.session_state["_cur_ld"] = st.session_state["ld_src"]

# ตัวแปรที่ใช้คำนวณ ให้โค้ดส่วน calc_losses มองเห็น
df_thk = st.session_state["thk_src"]
df_tdn = st.session_state["tdn_src"]
df_ld = st.session_state["ld_src"]

# ─────────────────────────────────────────────────────────────────────────────
# PRESTRESS LOSS ENGINE  (AASHTO LRFD 5.9.3)  — UNCHANGED
# ─────────────────────────────────────────────────────────────────────────────
def calc_losses(dft, dfp, fc, fci, fpu, fpi_ratio, aps_strand,
                num_tendon, n_strands, duct_dia_mm,
                t0, RH, anch_slip_mm, width):
    Ep    = 197_000.0
    mu    = 0.20
    K_wob = 0.0066
    KL    = 45.0
    b     = 1.0
    wc    = 2400.0

    x_mid  = width / 2.0
    t_mid  = float(np.interp(x_mid, dft["x (m)"], dft["t (m)"]))
    z_mid  = float(np.interp(x_mid, dfp["x (m)"], dfp["z_top (m)"]))
    yc_mid = t_mid / 2.0
    e_mid  = yc_mid - z_mid

    Ag_mid = b * t_mid
    Ig_mid = b * t_mid**3 / 12.0
    A_duct = math.pi / 4.0 * (duct_dia_mm / 1000.0)**2
    n_duct = int(num_tendon)
    y_duct = z_mid - yc_mid
    An_mid = Ag_mid - n_duct * A_duct
    In_mid = Ig_mid - n_duct * A_duct * y_duct**2
    VS     = (b * t_mid) / (2.0 * (b + t_mid))
    VS_mm  = VS * 1000.0

    aps_m2  = aps_strand * 1e-6
    n_total = int(num_tendon * n_strands)
    Aps     = n_total * aps_m2

    Ec  = 0.043 * (wc**1.5) * math.sqrt(fc)
    Eci = 0.043 * (wc**1.5) * math.sqrt(fci)

    xs = dft["x (m)"].values.astype(float)
    zs = np.interp(xs, dfp["x (m)"].values, dfp["z_top (m)"].values)
    dz = np.diff(zs); dx = np.diff(xs)
    dx_s = np.where(np.abs(dx) < 1e-9, 1e-9, dx)
    alpha = float(np.sum(np.abs(np.diff(np.append([0], np.arctan(dz / dx_s))))))
    alpha = max(alpha, 0.001)
    L_ten = float(np.sum(np.sqrt(dx**2 + dz**2)))
    if L_ten < 0.5: L_ten = float(xs[-1] - xs[0])

    fpj = fpu * fpi_ratio
    Pj  = Aps * fpj * 1e3

    exp_full = mu * alpha + K_wob * L_ten
    delta_fpF_full = fpj * (1.0 - math.exp(-exp_full))
    exp_mid  = mu * (alpha / 2.0) + K_wob * (L_ten / 2.0)
    delta_fpF = fpj * (1.0 - math.exp(-exp_mid))

    friction_slope = fpj * (mu * alpha / L_ten + K_wob)
    Lpa = math.sqrt((anch_slip_mm / 1000.0) * Ep / friction_slope)
    Lpa = min(Lpa, L_ten)
    delta_fpA = (anch_slip_mm / 1000.0) * Ep / Lpa
    delta_fpA = min(delta_fpA, 0.20 * fpj)

    fpt_mid = fpj - delta_fpF - (delta_fpA if Lpa > L_ten / 2.0 else 0.0)
    fpt_mid = max(fpt_mid, 0.5 * fpj)

    Pi_est = Aps * fpt_mid * 1e3
    fcgp   = (Pi_est/An_mid + Pi_est*e_mid**2/In_mid) / 1000.0
    delta_fpES = (Ep / Eci) * fcgp
    delta_fpES = max(0.0, delta_fpES)

    delta_imm   = delta_fpF + delta_fpA + delta_fpES
    imm_loss_pct = delta_imm / fpj * 100.0

    fpi_eff  = max(fpj - delta_imm, 0.5 * fpj)
    Pi_final = Aps * fpi_eff * 1e3

    fci_ksi = fci / 6.895
    kvs = max(1.45 - 0.0052 * VS_mm, 1.0)
    khs = max(2.00 - 0.014 * RH, 0.0)
    khc = max(1.56 - 0.008 * RH, 0.0)
    kf  = 5.0 / (1.0 + fci_ksi)
    ktd = 1.0

    eps_bdf  = kvs * khs * kf * ktd * 0.48e-3
    delta_fpSH = eps_bdf * Ep

    ti_safe = max(float(t0), 1.0)
    psi_b   = 1.9 * kvs * khc * kf * ktd * (ti_safe ** -0.118)

    fcgp_lt = (Pi_final/An_mid + Pi_final*e_mid**2/In_mid) / 1000.0
    delta_fpCR = max((Ep / Ec) * fcgp_lt * psi_b, 0.0)

    fpy = fpu * fpi_ratio / fpi_ratio * 0.9
    fpt_r = fpi_eff - 0.3 * (delta_fpSH + delta_fpCR)
    fpt_r = max(fpt_r, 0.5 * fpu)
    delta_fpR = max((fpt_r / KL) * (fpt_r / fpy - 0.55), 0.0)

    delta_lt    = delta_fpSH + delta_fpCR + delta_fpR
    lt_loss_pct = delta_lt / fpj * 100.0

    fpe_val  = max(fpi_eff - delta_lt, 0.45 * fpj)
    Pe_final = Aps * fpe_val * 1e3

    return dict(
        Ec=Ec, Eci=Eci, Ep=Ep,
        t_mid=t_mid, z_mid=z_mid, e_mid=e_mid,
        Ag_mid=Ag_mid, Ig_mid=Ig_mid, An_mid=An_mid, In_mid=In_mid,
        VS=VS, VS_mm=VS_mm,
        Aps=Aps, n_total=n_total,
        alpha=alpha, L_ten=L_ten, Lpa=Lpa, friction_slope=friction_slope,
        fpj=fpj, Pj=Pj,
        delta_fpF=delta_fpF, delta_fpF_full=delta_fpF_full,
        delta_fpA=delta_fpA, delta_fpES=delta_fpES,
        delta_imm=delta_imm, imm_loss_pct=imm_loss_pct,
        fcgp=fcgp, fpt_mid=fpt_mid, Pi=Pi_final,
        fci_ksi=fci_ksi, kvs=kvs, khs=khs, khc=khc, kf=kf, ktd=ktd,
        psi_b=psi_b, eps_bdf=eps_bdf,
        fcgp_lt=fcgp_lt, delta_fpSH=delta_fpSH,
        delta_fpCR=delta_fpCR, delta_fpR=delta_fpR,
        delta_lt=delta_lt, lt_loss_pct=lt_loss_pct,
        fpi_eff=fpi_eff, fpe=fpe_val, Pe=Pe_final,
        eff_ratio=Pe_final / Pj if Pj > 0 else 0.75,
        total_loss_pct=(delta_imm + delta_lt) / fpj * 100.0,
    )


# ─────────────────────────────────────────────────────────────────────────────
# 4.  CALCULATION ENGINE  — UNCHANGED
# ─────────────────────────────────────────────────────────────────────────────
def prep(df):
    df = df.copy()
    for col in df.columns:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna()
    if df.empty:
        return df
    return df.sort_values("x (m)").drop_duplicates(subset="x (m)").reset_index(drop=True)

def run_calc(dft, dfp, dfl, L):
    N = 500; b = 1.0
    x = np.linspace(0, width, N)

    t  = np.interp(x, dft["x (m)"], dft["t (m)"])
    z  = np.interp(x, dfp["x (m)"], dfp["z_top (m)"])
    yc = t / 2.0

    def ip(col): return np.interp(x, dfl["x (m)"], dfl[col])
    m_dl=ip("M_DL (kNm/m)"); v_dl=ip("V_DL (kN/m)")
    m_sdl=ip("M_SDL (kNm/m)"); v_sdl=ip("V_SDL (kN/m)")
    m_ll=ip("M_LL (kNm/m)");  v_ll=ip("V_LL (kN/m)")

    ms1 = m_dl + m_sdl + m_ll
    ms3 = m_dl + m_sdl + 0.8*m_ll
    mu  = 1.25*m_dl + 1.50*m_sdl + 1.75*m_ll
    vu  = 1.25*np.abs(v_dl) + 1.50*np.abs(v_sdl) + 1.75*np.abs(v_ll)

    Ag = b * t
    Ig = b * t**3 / 12.0

    A_duct = math.pi / 4.0 * (duct_dia_mm / 1000.0)**2
    n_ducts = int(num_tendon)
    y_duct  = z - yc
    An = Ag - n_ducts * A_duct
    In = Ig - n_ducts * A_duct * y_duct**2

    e = yc - z

    n_total = L["n_total"]
    Aps     = L["Aps"]
    fpi_val = L["fpi_eff"]
    Pi      = L["Pi"]
    Pe      = L["Pe"]

    def stress(P, M, ev, tv, Av, Iv):
        ht  = tv / 2.0
        top = (-P/Av + P*ev*ht/Iv - M*ht/Iv) / 1000.0
        bot = (-P/Av - P*ev*ht/Iv + M*ht/Iv) / 1000.0
        return top, bot

    tr_top,  tr_bot  = stress(Pi, m_dl, e, t, An, In)
    sv1_top, sv1_bot = stress(Pe, ms1,  e, t, Ag, Ig)
    sv3_top, sv3_bot = stress(Pe, ms3,  e, t, Ag, Ig)

    beta1 = float(np.clip(0.85 - 0.05*(fc-28.0)/7.0, 0.65, 0.85))
    k_fac = 2.0 * (1.04 - fpy_ratio)

    def flexure(dp_arr):
        dp_s = np.maximum(dp_arr, 1e-4)
        c_   = Aps*fpu / (0.85*fc*beta1*b*1000.0 + k_fac*Aps*fpu/dp_s)
        fps_ = np.clip(fpu*(1.0 - k_fac*c_/dp_s), 0.0, fpu)
        a_   = beta1 * c_
        Mn_  = Aps * fps_ * (dp_s - a_/2.0) * 1000.0
        return c_, a_, fps_, Mn_

    dp_pos = z;      dp_neg = t - z
    c_pos, a_pos, fps_pos, Mn_pos = flexure(dp_pos)
    c_neg, a_neg, fps_neg, Mn_neg = flexure(dp_neg)
    phi_Mn_pos =  phi_flex * Mn_pos
    phi_Mn_neg = -phi_flex * Mn_neg

    cdp_pos = np.where(dp_pos > 0, c_pos/dp_pos, np.inf)
    cdp_neg = np.where(dp_neg > 0, c_neg/dp_neg, np.inf)

    fr  = 0.63 * math.sqrt(fc)
    fpe = Pe / Ag / 1000.0
    Sb  = Ig / yc
    Mcr = (fr + fpe) * Sb / 1000.0

    dp_use = np.maximum(dp_pos, dp_neg)
    dv     = np.maximum(0.9*dp_use, 0.72*t)
    Vc     = 0.083*2.0*1.0*math.sqrt(fc)*b*dv*1000.0
    Vn_lim = 0.25*fc*b*dv*1000.0
    phi_Vn = phi_shear * np.minimum(Vc, Vn_lim)

    lim_tr_c  = -0.60*fci;  lim_tr_t  =  0.25*math.sqrt(fci)
    lim_sv_cp = -0.45*fc;   lim_sv_ct = -0.60*fc
    lim_sv_t  =  0.50*math.sqrt(fc)

    return dict(
        x=x, t=t, z=z, yc=yc, e=e,
        L=L,
        Ag=Ag, Ig=Ig, An=An, In=In, y_duct=y_duct,
        n_total=n_total, Aps=Aps, fpi_val=fpi_val, Pi=Pi, Pe=Pe,
        beta1=beta1, k_fac=k_fac,
        m_dl=m_dl, m_sdl=m_sdl, m_ll=m_ll,
        v_dl=v_dl, v_sdl=v_sdl, v_ll=v_ll,
        ms1=ms1, ms3=ms3, mu=mu, vu=vu,
        tr_top=tr_top, tr_bot=tr_bot,
        sv1_top=sv1_top, sv1_bot=sv1_bot,
        sv3_top=sv3_top, sv3_bot=sv3_bot,
        dp_pos=dp_pos, dp_neg=dp_neg,
        c_pos=c_pos, a_pos=a_pos, fps_pos=fps_pos,
        c_neg=c_neg, a_neg=a_neg, fps_neg=fps_neg,
        phi_Mn_pos=phi_Mn_pos, phi_Mn_neg=phi_Mn_neg,
        cdp_pos=cdp_pos, cdp_neg=cdp_neg,
        fr=fr, fpe=fpe, Sb=Sb, Mcr=Mcr,
        dv=dv, Vc=Vc, Vn_lim=Vn_lim, phi_Vn=phi_Vn,
        A_duct=A_duct, n_ducts=n_ducts,
        lim_tr_c=lim_tr_c, lim_tr_t=lim_tr_t,
        lim_sv_cp=lim_sv_cp, lim_sv_ct=lim_sv_ct, lim_sv_t=lim_sv_t,
    )


try:
    dft = prep(df_thk); dfp = prep(df_tdn); dfl = prep(df_ld)
    if any(len(d) < 2 for d in [dft, dfp, dfl]):
        st.warning("⚠️ Enter at least 2 rows in each table."); st.stop()

    t0_v   = int(st.session_state.get("t0", 3))
    rh_v   = int(st.session_state.get("RH", 75))
    anch_v = float(st.session_state.get("anch_slip_mm", 6.0))
    L = calc_losses(dft, dfp,
                    fc, fci, fpu, fpi_ratio, aps_strand,
                    num_tendon, n_strands, duct_dia_mm,
                    t0_v, rh_v, anch_v, width)
    R = run_calc(dft, dfp, dfl, L)

    sta_x   = dfl["x (m)"].values
    sta_idx = [int(np.abs(R["x"] - v).argmin()) for v in sta_x]
    N       = len(R["x"])

    # ─────────────────────────────────────────────────────────────────
    # SUMMARY METRIC CARDS
    # ─────────────────────────────────────────────────────────────────
    st.markdown("### 📊 Key Results")
    mc1, mc2, mc3, mc4, mc5 = st.columns(5)
    mc1.metric("Aps  (1m strip)",  f"{R['Aps']*1e6:.2f} mm²")
    mc2.metric("Pi  (initial)",    f"{R['Pi']:.2f} kN/m")
    mc3.metric("Pe  (effective)",  f"{R['Pe']:.2f} kN/m",
               delta=f"eff. ratio = {L['eff_ratio']:.3f}")
    mc4.metric("Total Loss",       f"{L['total_loss_pct']:.2f} %",
               delta=f"imm={L['imm_loss_pct']:.1f}%  lt={L['lt_loss_pct']:.1f}%")
    mc5.metric("Eff. Pe/Pi",       f"{L['eff_ratio']:.4f}")

    st.markdown("---")

    # ─────────────────────────────────────────────────────────────────
    # 5.  REPORT GENERATOR  — UNCHANGED logic
    # ─────────────────────────────────────────────────────────────────
    def make_report():
        doc = Document()
        for sec in doc.sections:
            sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
            sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        C_BLUE  = RGBColor(0x00, 0x44, 0x88)
        C_GREEN = RGBColor(0x00, 0x70, 0x00)
        C_RED   = RGBColor(0xC0, 0x00, 0x00)
        C_GRAY  = RGBColor(0x60, 0x60, 0x60)

        def h1(s): doc.add_heading(s, level=1)
        def h2(s): doc.add_heading(s, level=2)
        def h3(s): doc.add_heading(s, level=3)

        def para(s, bold=False, italic=False, color=None, indent=0.0, align=None):
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold=bold; r.italic=italic
            if color: r.font.color.rgb = color
            p.paragraph_format.left_indent = Inches(indent)
            if align: p.alignment = align
            return p

        def formula(s): return para(s, italic=True,  indent=0.5, color=C_GRAY)
        def subst(s):   return para(s, italic=True,  indent=0.7, color=C_GRAY)
        def result(s):  return para(s, bold=True,    indent=0.7, color=C_BLUE)
        def blank():    return doc.add_paragraph()

        def pf(cond, ok, fail):
            if cond: para(f"  ✔  {ok}   [PASS]",  bold=True, color=C_GREEN, indent=0.5)
            else:    para(f"  ✘  {fail}  [FAIL]",  bold=True, color=C_RED,   indent=0.5)

        def tbl(headers, rows, cw=None):
            t_ = doc.add_table(rows=1, cols=len(headers))
            t_.style = "Table Grid"
            for j,h in enumerate(headers):
                t_.rows[0].cells[j].text = h
                t_.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for row in rows:
                rc = t_.add_row().cells
                for j,v in enumerate(row): rc[j].text = str(v)
            if cw:
                for row in t_.rows:
                    for j,cell in enumerate(row.cells):
                        cell.width = Cm(cw[j])
            return t_

        def s(key, i): return float(R[key][i])

        # COVER
        blank(); blank()
        doc.add_heading("STRUCTURAL CALCULATION REPORT", 0)
        blank()
        tbl(["Item","Description"],[
            ["Project",       proj_name],
            ["Document No.",  doc_no],
            ["Subject",       "Transverse Tendon Design — PSC Box Girder Top Flange"],
            ["Code",          "AASHTO LRFD Bridge Design Specifications"],
            ["Prepared by",   eng_name],
            ["Checked by",    chk_name],
            ["Date",          datetime.datetime.now().strftime("%d %B %Y")],
        ], cw=[4.5,13.0])
        doc.add_page_break()

        # SEC 1 — DESIGN BASIS
        h1("1.  Design Basis")
        for it in [
            "Code: AASHTO LRFD Bridge Design Specifications",
            "Analysis basis: 1.0 m transverse strip across top flange",
            "Load combinations (AASHTO Table 3.4.1-1):",
            "  Strength I  :  1.25·DC + 1.50·DW + 1.75·LL",
            "  Service  I  :  1.00·DC + 1.00·DW + 1.00·LL  (compression check)",
            "  Service I   :  1.00·DC + 1.00·DW + 1.00·LL  (compression & tension check)",
            "  Transfer    :  Pi (after immediate losses) + M_DC",
            "Strand: Post-tensioned, bonded (fully grouted), low-relaxation",
            "Sign convention: Compression (−)  |  Tension (+)",
            "Positive moment = sagging (compression at TOP fibre)",
        ]: para(it, indent=0.3)
        blank()

        # SEC 2 — INPUT SUMMARY
        h1("2.  Design Input Summary")

        h2("2.1  Material Properties")
        tbl(["Parameter","Symbol","Value","Unit","Reference"],[
            ["Concrete — service",       "f'c",     f"{fc:.1f}",         "MPa","AASHTO 5.4.2"],
            ["Concrete — transfer",      "f'ci",    f"{fci:.1f}",        "MPa","AASHTO 5.9.2"],
            ["Strand tensile strength",  "fpu",     f"{fpu:.0f}",        "MPa","AASHTO 5.4.4"],
            ["Strand yield ratio",       "fpy/fpu", f"{fpy_ratio:.2f}",  "—",  "Low-relax"],
            ["Area per strand",          "asp",     f"{aps_strand:.1f}", "mm²","Product data"],
            ["PT duct outer diameter",   "d_duct",  f"{duct_dia_mm:.0f}","mm", "Supplier"],
        ], cw=[4.5,2.0,2.0,1.5,4.5])
        blank()

        h2("2.2  Prestressing Configuration")
        tbl(["Parameter","Symbol","Value","Unit"],[
            ["Tendons per 1 m strip",     "n_t",     f"{int(num_tendon)}",       "—"],
            ["Strands per tendon",        "n_s",     f"{int(n_strands)}",        "—"],
            ["Total strands (1m strip)",  "n",       f"{R['n_total']}",          "—"],
            ["Total Aps (1m strip)",      "Aps",     f"{R['Aps']*1e6:.2f}",     "mm²/m"],
            ["Jacking stress ratio",      "fpi/fpu", f"{fpi_ratio:.4f}",         "—"],
            ["Immediate loss (computed)", "Δfi",     f"{R['L']['imm_loss_pct']:.2f}", "%"],
            ["Long-term loss (computed)", "ΔfLT",    f"{R['L']['lt_loss_pct']:.2f}",  "%"],
            ["Total loss (computed)",     "Δftot",   f"{R['L']['total_loss_pct']:.2f}","%"],
        ], cw=[5.5,2.5,2.5,2.0])
        blank()

        h2("2.3  Resistance Factors")
        tbl(["Limit State","Symbol","Value"],[
            ["Flexure","φ_f",f"{phi_flex:.2f}"],
            ["Shear",  "φ_v",f"{phi_shear:.2f}"],
        ], cw=[6.0,2.5,2.5])
        blank()

        h2("2.4  Allowable Stress Limits")
        tbl(["Condition","Expression","Limit (MPa)","Article"],[
            ["Transfer — Compression",         "−0.60·f'ci", f"{R['lim_tr_c']:.3f}","5.9.2.3.1a"],
            ["Transfer — Tension (bonded)",    "+0.62·√f'ci",f"+{R['lim_tr_t']:.4f}","5.9.2.3.1b"],
            ["Service I — Comp (perm.loads)",  "−0.45·f'c",  f"{R['lim_sv_cp']:.3f}","5.9.2.3.2a"],
            ["Service I — Comp (total loads)", "−0.60·f'c",  f"{R['lim_sv_ct']:.3f}","5.9.2.3.2a"],
            ["Service I — Tension (bonded)",   "+0.50·√f'c", f"+{R['lim_sv_t']:.4f}","5.9.2.3.2b"],
        ], cw=[5.5,3.5,2.5,2.5])
        blank()

        h2("2.5  Input Geometry and Load at Stations")
        geo_rows = []
        for i in sta_idx:
            geo_rows.append([
                f"{R['x'][i]:.2f}",
                f"{R['t'][i]*1000:.2f}", f"{R['z'][i]*1000:.2f}", f"{R['yc'][i]*1000:.2f}",
                f"{R['e'][i]*1000:.2f}",
                f"{R['m_dl'][i]:.2f}",   f"{R['m_sdl'][i]:.2f}",  f"{R['m_ll'][i]:.2f}",
                f"{R['v_dl'][i]:.2f}",   f"{R['v_sdl'][i]:.2f}",  f"{R['v_ll'][i]:.2f}",
            ])
        tbl(["x(m)","t(mm)","z(mm)","yc(mm)","e(mm)",
             "M_DL","M_SDL","M_LL","V_DL","V_SDL","V_LL"],
            geo_rows, cw=[1.4,1.4,1.4,1.4,1.4,1.6,1.6,1.6,1.6,1.6,1.6])
        para("  M in kNm/m  |  V in kN/m", italic=True, color=C_GRAY)
        doc.add_page_break()

        # SEC 3 — PRESTRESS LOSSES
        h1("3.  Prestress Loss Calculation  (AASHTO LRFD 5.9.3)")
        _L = R["L"]
        _fpj = _L["fpj"]

        h2("3.1  Material Properties")
        formula("Ec  =  0.043 × wc^1.5 × √f'c   [AASHTO 5.4.2.4]  (wc = 2400 kg/m³)")
        subst( f"    =  0.043 × 2400^1.5 × √{fc:.1f}")
        result(f"    =  {_L['Ec']:.2f} MPa")
        blank()
        formula("Eci =  0.043 × wc^1.5 × √f'ci")
        subst( f"    =  0.043 × 2400^1.5 × √{fci:.1f}")
        result(f"    =  {_L['Eci']:.2f} MPa")
        blank()
        para("Ep = 197,000 MPa  (AASHTO 5.4.4.2, standard)", indent=0.3)
        blank()

        h2("3.2  Section Properties at Midspan  (x = L/2)")
        tbl(["Property","Value","Unit"],[
            ["Slab thickness",            f"{_L['t_mid']*1000:.2f}", "mm"],
            ["Tendon CG from top",         f"{_L['z_mid']*1000:.2f}", "mm"],
            ["Eccentricity e = yc−z",      f"{_L['e_mid']*1000:.2f}", "mm"],
            ["Gross area Ag",              f"{_L['Ag_mid']*1e6:.2f}", "mm²/m"],
            ["Net area An (duct deducted)",f"{_L['An_mid']*1e6:.2f}", "mm²/m"],
            ["Volume/Surface ratio",       f"{_L['VS_mm']:.2f}", "mm"],
        ], cw=[6,4,2])
        blank()

        h2("3.3  Tendon Geometry")
        formula("Angular change α:  computed from ∑|Δθ| along tendon profile")
        result(f"  α  =  {_L['alpha']:.6f} rad")
        blank()
        formula("Tendon length:  L = ∑√(Δx²+Δz²)")
        result(f"  L  =  {_L['L_ten']:.4f} m")
        blank()
        para(f"Jacking method: ONE-END  |  fpj = 0.75×fpu = 0.75×{fpu:.0f} = {_fpj:.2f} MPa",
             indent=0.3)
        blank()

        h2("3.4  Immediate Losses")

        h3("3.4.1  Friction Loss  ΔfpF  (AASHTO 5.9.3.2.1)")
        para("Parameters: μ = 0.20, K = 0.0066 rad/m  (7-wire low-relax, internal grouted)",
             italic=True, indent=0.3)
        formula("ΔfpF  =  fpj × (1 − e^(−μα − Kx))")
        subst( f"      =  {_fpj:.2f} × (1 − e^(−{0.20:.2f}×{_L['alpha']:.4f} − {0.0066:.4f}×{_L['L_ten']/2:.3f}))")
        result(f"ΔfpF  =  {_L['delta_fpF']:.4f} MPa  ({_L['delta_fpF']/_fpj*100:.2f}% of fpj)")
        blank()

        h3("3.4.2  Anchorage Set Loss  ΔfpA  (AASHTO 5.9.3.2.2)")
        anch_s = float(st.session_state.get("anch_slip_mm", 6.0))
        formula("Length affected by anchor set:")
        formula("  Lpa  =  √[ Δ·Ep / (μ·fpj/L + K) / 1000 ]")
        subst( f"       =  √[ {anch_s:.1f}mm×197000 / ({0.20:.2f}×{_fpj:.2f}/{_L['L_ten']:.3f}+{0.0066:.4f}) / 1000 ]")
        result(f"  Lpa  =  {_L['Lpa']:.4f} m")
        blank()
        formula("ΔfpA  =  Δ × Ep / Lpa")
        subst( f"      =  {anch_s:.1f}mm × 197,000 MPa / {_L['Lpa']*1000:.2f}mm")
        result(f"ΔfpA  =  {_L['delta_fpA']:.4f} MPa  ({_L['delta_fpA']/_fpj*100:.2f}% of fpj)")
        blank()

        h3("3.4.3  Elastic Shortening  ΔfpES  (AASHTO 5.9.3.2.3)")
        formula("fcgp  =  (Pi/An + Pi·e²/In) / 1000   [concrete stress at tendon CG]")
        Pi_v = float(_L["Pi"]); e_v = float(_L["e_mid"])
        An_v = float(_L["An_mid"]); In_v = float(_L["In_mid"])
        subst( f"      =  ({Pi_v:.4f}/{An_v*1e6:.2f}mm²  +  {Pi_v:.4f}×{e_v*1000:.2f}²mm²/{In_v*1e12:.4f}×10⁻³mm⁴) / 1000")
        result(f"fcgp  =  {_L['fcgp']:.4f} MPa")
        blank()
        formula("ΔfpES  =  (Ep / Eci) × fcgp")
        subst( f"       =  (197,000 / {_L['Eci']:.2f}) × {_L['fcgp']:.4f}")
        result(f"ΔfpES  =  {_L['delta_fpES']:.4f} MPa  ({_L['delta_fpES']/_fpj*100:.2f}% of fpj)")
        blank()

        h3("3.4.4  Total Immediate Loss")
        formula("Δfi  =  ΔfpF + ΔfpA + ΔfpES")
        subst( f"     =  {_L['delta_fpF']:.4f} + {_L['delta_fpA']:.4f} + {_L['delta_fpES']:.4f}")
        result(f"Δfi  =  {_L['delta_imm']:.4f} MPa  ({_L['imm_loss_pct']:.2f}% of fpj)")
        blank()
        formula("fpi (effective)  =  fpj − Δfi")
        subst( f"                 =  {_fpj:.4f} − {_L['delta_imm']:.4f}")
        result(f"fpi  =  {_L['fpi_eff']:.4f} MPa")
        blank()
        formula("Pi  =  Aps × fpi")
        subst( f"    =  {_L['Aps']*1e6:.2f} mm²  ×  {_L['fpi_eff']:.4f} MPa  × 10⁻³")
        result(f"Pi  =  {_L['Pi']:.4f} kN/m")
        blank()

        h2("3.5  Long-Term Losses  (Approximate Method, AASHTO 5.9.3.3)")
        para(f"t₀ = {t0_v} days  |  RH = {rh_v}%  |  V/S = {_L['VS_mm']:.1f} mm",
             italic=True, indent=0.3)
        blank()

        h3("3.5.1  Correction Factors")
        formula("kvs  =  max(1.45 − 0.0052·V/S,  1.0)")
        subst( f"     =  max(1.45 − 0.0052×{_L['VS_mm']:.2f},  1.0)")
        result(f"kvs  =  {_L['kvs']:.4f}")
        blank()
        formula("khs  =  2.00 − 0.014·RH   (shrinkage humidity)")
        subst( f"     =  2.00 − 0.014×{rh_v}")
        result(f"khs  =  {_L['khs']:.4f}")
        blank()
        formula("khc  =  1.56 − 0.008·RH   (creep humidity)")
        subst( f"     =  1.56 − 0.008×{rh_v}")
        result(f"khc  =  {_L['khc']:.4f}")
        blank()
        formula("kf   =  5.0 / (1 + f'ci)")
        subst( f"     =  5.0 / (1 + {fci:.1f})")
        result(f"kf   =  {_L['kf']:.4f}")
        blank()
        formula("γh  (humidity factor for creep/shrinkage) = 1.7 − 0.01·RH")
        subst( f"     =  1.7 − 0.01×{rh_v}")
        result(f"γh   =  {1.7 - 0.01*float(st.session_state.get('RH',75)):.4f}  (= 1.7 − 0.01×RH)")
        blank()
        formula("γst  (concrete strength factor) = 5.0 / (1 + f'ci_ksi)")
        result(f"γst  =  {5.0/(1.0+fci/6.895):.4f}  (= 5/(1+f'ci_ksi))")
        blank()
        formula("ψb (creep coeff.)  =  1.9·kvs·khc·kf·ktd·t₀^(−0.118)  [5.4.2.3.2]")
        subst( f"                   =  1.9×{_L['kvs']:.4f}×{_L['khc']:.4f}×{_L['kf']:.4f}×1.0×{t0_v}^(−0.118)")
        result(f"ψb  =  {_L['psi_b']:.4f}")
        blank()

        h3("3.5.2  Creep Loss  ΔfpCR  (AASHTO 5.9.3.3-1)")
        formula("ΔfpCR  =  10.0 × (fpi·Aps/Ag) × γh × γst")
        subst( f"       =  ({197000:.0f}/{_L['Ec']:.0f}) × {_L['fcgp_lt']:.4f} × {_L['psi_b']:.4f}")
        result(f"ΔfpCR  =  {_L['delta_fpCR']:.4f} MPa")
        blank()

        h3("3.5.3  Shrinkage Loss  ΔfpSH  (AASHTO 5.9.3.3-1)")
        formula("ΔfpSH  =  83 × γh × γst")
        subst( f"       =  {_L['eps_bdf']:.6f} × {197000:.0f}")
        result(f"ΔfpSH  =  {_L['delta_fpSH']:.4f} MPa")
        blank()

        h3("3.5.4  Relaxation Loss  ΔfpR  (AASHTO 5.9.3.4.3, low-relax)")
        formula("ΔfpR   =  0.3 × [20.0 − 0.4·ΔfpES − 0.2·(ΔfpCR + ΔfpSH)]")
        subst( f"       =  0.3 × [20.0 − 0.4×{_L['delta_fpES']:.4f} − 0.2×({_L['delta_fpCR']:.4f}+{_L['delta_fpSH']:.4f})]")
        result(f"ΔfpR   =  {_L['delta_fpR']:.4f} MPa")
        blank()

        h3("3.5.5  Total Long-Term Loss & Effective Prestress")
        formula("ΔfLT   =  ΔfpCR + ΔfpSH + ΔfpR")
        subst( f"       =  {_L['delta_fpCR']:.4f} + {_L['delta_fpSH']:.4f} + {_L['delta_fpR']:.4f}")
        result(f"ΔfLT   =  {_L['delta_lt']:.4f} MPa  ({_L['lt_loss_pct']:.2f}% of fpi)")
        blank()
        formula("Δftotal  =  Δfi + ΔfLT")
        subst( f"         =  {_L['delta_imm']:.4f} + {_L['delta_lt']:.4f}")
        result(f"Δftotal  =  {_L['delta_imm']+_L['delta_lt']:.4f} MPa  ({_L['total_loss_pct']:.2f}% of fpj)")
        blank()
        formula("fpe  =  fpj − Δftotal")
        subst( f"     =  {_fpj:.4f} − {_L['delta_imm']+_L['delta_lt']:.4f}")
        result(f"fpe  =  {_L['fpe']:.4f} MPa")
        blank()
        formula("Pe   =  Aps × fpe")
        subst( f"     =  {_L['Aps']*1e6:.2f} mm²  ×  {_L['fpe']:.4f} MPa  × 10⁻³")
        result(f"Pe   =  {_L['Pe']:.4f} kN/m")
        doc.add_page_break()

        h1("4.  Global Prestress Force Summary")

        h2("4.1  Total Prestress Steel Area  Aps")
        formula("Aps  =  n_total × asp")
        subst( f"     =  {R['n_total']} strands  ×  {aps_strand:.1f} mm²/strand")
        result(f"     =  {R['Aps']*1e6:.4f} mm²/m")
        blank()

        h2("4.2  Jacking Stress  fpi  (after immediate losses)")
        formula("fpi  =  fpu × (fpi/fpu) × (1 − Δi/100)")
        subst( f"     =  {fpu:.0f} × {fpi_ratio:.4f} × (1 − {_L['imm_loss_pct']:.1f}/100)")
        result(f"     =  {R['fpi_val']:.4f} MPa")
        blank()

        h2("4.3  Initial Prestress Force  Pi")
        formula("Pi   =  Aps × fpi  × 10⁻³")
        subst( f"     =  {R['Aps']*1e6:.4f} mm²/m  ×  {R['fpi_val']:.4f} MPa  × 10⁻³")
        result(f"     =  {R['Pi']:.4f} kN/m")
        blank()

        h2("4.4  Effective Prestress Force  Pe  (after all losses)")
        formula("Pe   =  Pi × (Pe/Pi)")
        subst( f"     =  {R['Pi']:.4f}  ×  {_L['eff_ratio']:.4f}")
        result(f"     =  {R['Pe']:.4f} kN/m")

        h2("4.5  Section Factors")
        formula("β₁  =  0.85 − 0.05 × (f'c − 28.0)/7.0   [0.65 ≤ β₁ ≤ 0.85]")
        subst( f"    =  0.85 − 0.05 × ({fc:.1f} − 28.0)/7.0")
        result(f"    =  {R['beta1']:.4f}")
        blank()
        formula("k   =  2.0 × (1.04 − fpy/fpu)   [AASHTO C5.6.3.1.1]")
        subst( f"    =  2.0 × (1.04 − {fpy_ratio:.2f})")
        result(f"    =  {R['k_fac']:.4f}")
        doc.add_page_break()

        # SEC 5 — STATION-BY-STATION
        h1("5.  Detailed Station-by-Station Calculations")
        para("Calculations are presented per 1.0 m strip width at each station.", italic=True)
        blank()

        for ks, i in enumerate(sta_idx):
            xi   = float(R["x"][i])
            ti   = float(R["t"][i]);       zi   = float(R["z"][i])
            yci  = float(R["yc"][i]);      ei   = float(R["e"][i])
            Agi  = float(R["Ag"][i]);      Igi  = float(R["Ig"][i])
            Ani  = float(R["An"][i]);      Ini  = float(R["In"][i])
            ydi  = float(R["y_duct"][i])
            mdi  = float(R["m_dl"][i]);    msdi = float(R["m_sdl"][i])
            mli  = float(R["m_ll"][i]);    vdi  = float(R["v_dl"][i])
            vsdi = float(R["v_sdl"][i]);   vli  = float(R["v_ll"][i])
            ms1i = float(R["ms1"][i]);     ms3i = float(R["ms3"][i])
            mui  = float(R["mu"][i]);      vui  = float(R["vu"][i])
            trt  = float(R["tr_top"][i]);  trb  = float(R["tr_bot"][i])
            s1t  = float(R["sv1_top"][i]); s1b  = float(R["sv1_bot"][i])
            dpp  = float(R["dp_pos"][i]);  dpn  = float(R["dp_neg"][i])
            cpp  = float(R["c_pos"][i]);   app  = float(R["a_pos"][i])
            fpp  = float(R["fps_pos"][i])
            cpn  = float(R["c_neg"][i]);   apn  = float(R["a_neg"][i])
            fpn  = float(R["fps_neg"][i])
            pMp  = float(R["phi_Mn_pos"][i])
            pMn_ = float(R["phi_Mn_neg"][i])
            cdpp = float(R["cdp_pos"][i]); cdpn = float(R["cdp_neg"][i])
            fpei = float(R["fpe"][i]);     Sbi  = float(R["Sb"][i])
            Mcri = float(R["Mcr"][i])
            fri  = float(R["fr"])
            dvi  = float(R["dv"][i]);      Vci  = float(R["Vc"][i])
            pVi  = float(R["phi_Vn"][i]); Vnli = float(R["Vn_lim"][i])
            A_d  = float(R["A_duct"])
            n_d  = int(R["n_ducts"])

            ltr_c = float(R["lim_tr_c"]);  ltr_t = float(R["lim_tr_t"])
            lsv_ct= float(R["lim_sv_ct"]); lsv_t = float(R["lim_sv_t"])

            doc.add_heading(f"4.{ks+1}   Station  x = {xi:.2f} m", level=2)

            h3(f"5.{ks+1}.1   Net Section Properties  (duct deducted — used at Transfer)")
            tbl(["Property","Formula","Substitution","Value","Unit"],[
                ["Slab thickness",      "t",          "input",                f"{ti*1000:.2f}","mm"],
                ["Tendon CG from top",  "z",          "input",                f"{zi*1000:.2f}","mm"],
                ["Section centroid",    "yc = t/2",   f"{ti*1000:.2f}/2",     f"{yci*1000:.2f}","mm"],
                ["Eccentricity",        "e = yc − z", f"{yci*1000:.2f}−{zi*1000:.2f}", f"{ei*1000:.4f}","mm"],
                ["Gross area",          "Ag = 1000·t",f"1000×{ti*1000:.2f}",  f"{Agi*1e6:.2f}","mm²/m"],
                ["Gross inertia",       "Ig = 1000·t³/12",f"1000×{ti*1000:.2f}³/12",f"{Igi*1e12:.4f}×10⁻³","mm⁴/m"],
                ["Duct area (each)",    "Ad = π·d²/4",f"π×{duct_dia_mm:.0f}²/4",f"{A_d*1e6:.3f}","mm²"],
                ["Duct CG from CG",     "yd = z−yc",  f"{zi*1000:.2f}−{yci*1000:.2f}",f"{ydi*1000:.4f}","mm"],
                ["Net area",            "An = Ag − n·Ad",f"{Agi*1e6:.2f}−{n_d}×{A_d*1e6:.3f}",f"{Ani*1e6:.4f}","mm²/m"],
                ["Net inertia",         "In = Ig − n·Ad·yd²",
                 f"{Igi*1e12:.4f}×10⁻³−{n_d}×{A_d*1e6:.3f}×{ydi*1000:.4f}²×10⁻⁶",
                 f"{Ini*1e12:.6f}×10⁻³","mm⁴/m"],
            ], cw=[3.5,3.5,5.5,2.5,1.5])
            blank()

            h3(f"5.{ks+1}.2   Load Combinations")
            tbl(["Combination","Expression","Substitution","Value","Unit"],[
                ["Service I",
                 "Ms1 = M_DL + M_SDL + M_LL",
                 f"{mdi:.2f}+{msdi:.2f}+{mli:.2f}", f"{ms1i:.4f}","kNm/m"],
                ["Strength I — Moment",
                 "Mu = 1.25·MDL + 1.50·MSDL + 1.75·MLL",
                 f"1.25×{mdi:.2f}+1.50×{msdi:.2f}+1.75×{mli:.2f}",
                 f"{mui:.4f}","kNm/m"],
                ["Strength I — Shear",
                 "Vu = 1.25|VDL| + 1.50|VSDL| + 1.75|VLL|",
                 f"1.25×|{vdi:.2f}|+1.50×|{vsdi:.2f}|+1.75×|{vli:.2f}|",
                 f"{vui:.4f}","kN/m"],
            ], cw=[2.5,5.0,5.0,2.0,1.5])
            blank()

            h3(f"5.{ks+1}.3   Stress Check — Transfer  (AASHTO 5.9.2.3.1)")
            para("Loading: Pi + M_DL  |  Net section (duct deducted)", italic=True, indent=0.3)
            blank()
            para("Stress formula:", bold=True, indent=0.3)
            formula("σ_top = [ −Pi/An  +  Pi·e·yc/In  −  M·yc/In ] × 10⁻³  (MPa)")
            formula("σ_bot = [ −Pi/An  −  Pi·e·yc/In  +  M·yc/In ] × 10⁻³  (MPa)")
            blank()
            para("TOP fibre:", bold=True, indent=0.3)
            formula(f"σ_tr,top = [−{R['Pi']:.4f}/{Ani*1e6:.4f}"
                    f" + {R['Pi']:.4f}×{ei*1000:.4f}×{yci*1000:.4f}/{Ini*1e12:.6f}×10⁻³"
                    f" − {mdi:.4f}×{yci*1000:.4f}/{Ini*1e12:.6f}×10⁻³] × 10⁻³")
            result(f"σ_tr,top  =  {trt:.6f} MPa")
            pf(ltr_c <= trt <= ltr_t,
               f"σ_tr,top = {trt:.4f} MPa  within [{ltr_c:.3f},  +{ltr_t:.4f}] MPa",
               f"σ_tr,top = {trt:.4f} MPa  outside [{ltr_c:.3f}, +{ltr_t:.4f}] MPa")
            blank()
            para("BOTTOM fibre:", bold=True, indent=0.3)
            formula(f"σ_tr,bot = [−{R['Pi']:.4f}/{Ani*1e6:.4f}"
                    f" − {R['Pi']:.4f}×{ei*1000:.4f}×{yci*1000:.4f}/{Ini*1e12:.6f}×10⁻³"
                    f" + {mdi:.4f}×{yci*1000:.4f}/{Ini*1e12:.6f}×10⁻³] × 10⁻³")
            result(f"σ_tr,bot  =  {trb:.6f} MPa")
            pf(ltr_c <= trb <= ltr_t,
               f"σ_tr,bot = {trb:.4f} MPa  within [{ltr_c:.3f},  +{ltr_t:.4f}] MPa",
               f"σ_tr,bot = {trb:.4f} MPa  outside [{ltr_c:.3f}, +{ltr_t:.4f}] MPa")
            blank()

            h3(f"5.{ks+1}.4   Stress Check — Service I  (AASHTO 5.9.2.3.2)")
            para("Gross section used (ducts grouted).  Loading: Pe + load combination.",
                 italic=True, indent=0.3)
            blank()
            for (combo_name, M_i, t_s, b_s) in [
                ("Service I  (compression & tension check)", ms1i, s1t, s1b),
            ]:
                para(f"── {combo_name}  |  M = {M_i:.4f} kNm/m ──", bold=True, indent=0.3)
                formula(f"σ_top = [−{R['Pe']:.4f}/{Agi*1e6:.4f}"
                        f" + {R['Pe']:.4f}×{ei*1000:.4f}×{yci*1000:.4f}/{Igi*1e12:.6f}×10⁻³"
                        f" − {M_i:.4f}×{yci*1000:.4f}/{Igi*1e12:.6f}×10⁻³] × 10⁻³")
                result(f"σ_top  =  {t_s:.6f} MPa")
                pf(t_s >= lsv_ct,
                   f"σ_top = {t_s:.4f} MPa  ≥  {lsv_ct:.3f} MPa  (−0.60·f'c)",
                   f"σ_top = {t_s:.4f} MPa  <   {lsv_ct:.3f} MPa  EXCEEDS LIMIT")
                blank()
                formula(f"σ_bot = [−{R['Pe']:.4f}/{Agi*1e6:.4f}"
                        f" − {R['Pe']:.4f}×{ei*1000:.4f}×{yci*1000:.4f}/{Igi*1e12:.6f}×10⁻³"
                        f" + {M_i:.4f}×{yci*1000:.4f}/{Igi*1e12:.6f}×10⁻³] × 10⁻³")
                result(f"σ_bot  =  {b_s:.6f} MPa")
                pf(b_s >= lsv_ct,
                   f"σ_bot = {b_s:.4f} MPa  ≥  {lsv_ct:.3f} MPa  (−0.60·f'c)",
                   f"σ_bot = {b_s:.4f} MPa  <   {lsv_ct:.3f} MPa  EXCEEDS LIMIT")
                pf(t_s <= lsv_t,
                   f"σ_top = {t_s:.4f} MPa  ≤  +{lsv_t:.4f} MPa  (tension OK)",
                   f"σ_top = {t_s:.4f} MPa  >  +{lsv_t:.4f} MPa  TENSION EXCEEDED")
                pf(b_s <= lsv_t,
                   f"σ_bot = {b_s:.4f} MPa  ≤  +{lsv_t:.4f} MPa  (tension OK)",
                   f"σ_bot = {b_s:.4f} MPa  >  +{lsv_t:.4f} MPa  TENSION EXCEEDED")
                blank()

            h3(f"5.{ks+1}.5   Flexural Strength Check — Strength I  (AASHTO 5.6.3)")
            para("Rectangular stress block | No mild steel | Separate +Mu / −Mu capacity",
                 italic=True, indent=0.3)
            blank()
            for (label, dp_v, c_v, a_v, fp_v, pMnv, cdpv, mux) in [
                ("+Mu  (sagging, comp. face = TOP)",       dpp, cpp, app, fpp,  pMp,       cdpp, mui),
                ("−Mu  (hogging, comp. face = BOTTOM)",   dpn, cpn, apn, fpn,  abs(pMn_), cdpn, mui),
            ]:
                para(f"── {label} ──", bold=True, indent=0.3)
                para(f"  Effective depth  dp = {dp_v*1000:.2f} mm", indent=0.4)
                blank()
                para("  Step 1  Depth of neutral axis  c:", bold=True, indent=0.3)
                formula("  c  =  Aps·fpu / (0.85·f'c·β₁·b·1000  +  k·Aps·fpu / dp)")
                subst (f"     =  {R['Aps']*1e6:.4f}×{fpu:.0f}"
                       f" / (0.85×{fc:.1f}×{R['beta1']:.4f}×1000"
                       f" + {R['k_fac']:.4f}×{R['Aps']*1e6:.4f}×{fpu:.0f}/{dp_v*1000:.2f})")
                result(f"  c  =  {c_v*1000:.4f} mm")
                blank()
                para("  Step 2  Depth of stress block  a  =  β₁·c:", bold=True, indent=0.3)
                formula(f"  a  =  {R['beta1']:.4f}  ×  {c_v*1000:.4f} mm")
                result(f"  a  =  {a_v*1000:.4f} mm")
                pf(a_v <= dp_v,
                   f"a ({a_v*1000:.2f} mm) ≤ dp ({dp_v*1000:.2f} mm)  — rectangular section OK",
                   f"a ({a_v*1000:.2f} mm) > dp ({dp_v*1000:.2f} mm)  — T-section!")
                blank()
                para("  Step 3  Stress in prestress steel  fps:", bold=True, indent=0.3)
                formula("  fps  =  fpu × [1 − k·(c/dp)]")
                subst (f"      =  {fpu:.0f} × [1 − {R['k_fac']:.4f}×{c_v*1000:.4f}/{dp_v*1000:.2f}]")
                result(f"  fps  =  {fp_v:.4f} MPa")
                blank()
                para("  Step 4  Nominal flexural resistance  Mn:", bold=True, indent=0.3)
                formula("  Mn   =  Aps · fps · (dp − a/2)")
                subst (f"      =  {R['Aps']*1e6:.4f}mm²  ×  {fp_v:.4f}MPa"
                       f"  ×  ({dp_v*1000:.2f} − {a_v*1000:.4f}/2)mm  × 10⁻⁶")
                result(f"  Mn   =  {pMnv/phi_flex:.4f} kNm/m")
                blank()
                para("  Step 5  Factored resistance  φMn:", bold=True, indent=0.3)
                formula(f"  φMn  =  {phi_flex:.2f}  ×  {pMnv/phi_flex:.4f}")
                result(f"  φMn  =  {pMnv:.4f} kNm/m")
                blank()
                para("  Step 6  Demand/Capacity  (DCR):", bold=True, indent=0.3)
                dcr_v = abs(mux)/pMnv if pMnv > 0 else 999
                pf(abs(mux) <= pMnv,
                   f"|Mu|={abs(mux):.4f} ≤ φMn={pMnv:.4f} kNm/m  (DCR={dcr_v:.4f})",
                   f"|Mu|={abs(mux):.4f} > φMn={pMnv:.4f} kNm/m  (DCR={dcr_v:.4f})  FAILS")
                blank()
                para("  Step 7  Ductility  c/dp ≤ 0.42  (AASHTO 5.7.3.3.1):", bold=True, indent=0.3)
                formula(f"  c/dp  =  {c_v*1000:.4f} / {dp_v*1000:.2f}  =  {cdpv:.4f}")
                pf(cdpv <= 0.42,
                   f"c/dp = {cdpv:.4f} ≤ 0.42  — tension-controlled",
                   f"c/dp = {cdpv:.4f} > 0.42  — NOT tension-controlled")
                blank()

            para("── Minimum Reinforcement  (AASHTO 5.6.3.3) ──", bold=True, indent=0.3)
            formula("Mcr  =  (fr + fpe) × Sb  × 10⁻³")
            formula(f"     =  ({fri:.4f} MPa  +  {fpei:.4f} MPa)  ×  {Sbi:.8f} m³")
            result(f"Mcr  =  {Mcri:.4f} kNm/m")
            blank()
            min_req = min(1.2*Mcri, 1.33*abs(mui))
            formula(f"1.2·Mcr = {1.2*Mcri:.4f} kNm/m")
            formula(f"1.33·|Mu| = {1.33*abs(mui):.4f} kNm/m   →  governing = {min_req:.4f} kNm/m")
            pf(max(pMp, abs(pMn_)) >= min_req,
               f"φMn = {max(pMp, abs(pMn_)):.4f} ≥ {min_req:.4f} kNm/m  OK",
               f"φMn = {max(pMp, abs(pMn_)):.4f} < {min_req:.4f} kNm/m  INSUFFICIENT")
            blank()

            h3(f"5.{ks+1}.6   Shear Strength Check — Strength I  (AASHTO 5.7.3)")
            para("Simplified method: β=2.0  |  Vs=0 (no stirrups)  |  Vp=0",
                 italic=True, indent=0.3)
            blank()
            para("  Step 1  Effective shear depth  dv  (AASHTO 5.7.2.8):", bold=True, indent=0.3)
            dp_use_v = max(dpp, dpn)
            formula("  dv  =  max(0.9·dp,  0.72·t)")
            subst (f"      =  max(0.9×{dp_use_v*1000:.2f}mm,  0.72×{ti*1000:.2f}mm)")
            result(f"  dv  =  {dvi*1000:.4f} mm")
            blank()
            para("  Step 2  Concrete shear resistance  Vc  (AASHTO 5.7.3.3-3):", bold=True, indent=0.3)
            formula("  Vc  =  0.083·β·λ·√f'c·bv·dv × 10⁻³")
            subst (f"      =  0.083×2.0×1.0×√{fc:.1f}×1000mm×{dvi*1000:.4f}mm × 10⁻³")
            result(f"  Vc  =  {Vci:.4f} kN/m")
            blank()
            para("  Step 3  Upper limit  Vn,max  (AASHTO 5.7.3.3-2):", bold=True, indent=0.3)
            formula("  Vn,max  =  0.25·f'c·bv·dv × 10⁻³")
            subst (f"         =  0.25×{fc:.1f}MPa×1000mm×{dvi*1000:.4f}mm × 10⁻³")
            result(f"  Vn,max  =  {Vnli:.4f} kN/m")
            blank()
            Vn_use = min(Vci, Vnli)
            para("  Step 4  Nominal shear resistance:", bold=True, indent=0.3)
            formula("  Vn  =  min(Vc, Vn,max)  [Vs=0, Vp=0]")
            result(f"  Vn  =  {Vn_use:.4f} kN/m")
            blank()
            para("  Step 5  Factored resistance  φVn:", bold=True, indent=0.3)
            formula(f"  φVn  =  {phi_shear:.2f}  ×  {Vn_use:.4f}")
            result(f"  φVn  =  {pVi:.4f} kN/m")
            blank()
            para("  Step 6  Demand/Capacity check:", bold=True, indent=0.3)
            dcr_sh = vui/pVi if pVi > 0 else 999
            pf(vui <= pVi,
               f"Vu={vui:.4f} ≤ φVn={pVi:.4f} kN/m  (DCR={dcr_sh:.4f})",
               f"Vu={vui:.4f} > φVn={pVi:.4f} kN/m  (DCR={dcr_sh:.4f})  INSUFFICIENT")
            blank()
            doc.add_page_break()

        # SEC 6 — SUMMARY
        h1("6.  Summary of Results — All Stations")
        sum_rows = []
        for i in sta_idx:
            mui_ = float(R["mu"][i]); vui_ = float(R["vu"][i])
            pMp_ = float(R["phi_Mn_pos"][i]); pMn__ = float(R["phi_Mn_neg"][i])
            pVi_ = float(R["phi_Vn"][i])
            cap  = pMp_ if mui_ >= 0 else abs(pMn__)
            dcr_m = abs(mui_)/cap   if cap  > 0 else 999
            dcr_v = vui_/pVi_       if pVi_ > 0 else 999
            ok_tr = (R["lim_tr_c"]<=R["tr_top"][i]<=R["lim_tr_t"] and
                     R["lim_tr_c"]<=R["tr_bot"][i]<=R["lim_tr_t"])
            ok_sv = (R["sv1_top"][i] >= R["lim_sv_ct"] and
                     R["sv1_bot"][i] >= R["lim_sv_ct"] and
                     R["sv1_top"][i] <= R["lim_sv_t"]  and
                     R["sv1_bot"][i] <= R["lim_sv_t"])
            sum_rows.append([
                f"{R['x'][i]:.2f}",
                f"{R['tr_top'][i]:.3f}",  f"{R['tr_bot'][i]:.3f}",
                "PASS" if ok_tr else "FAIL",
                f"{R['sv1_top'][i]:.3f}", f"{R['sv1_bot'][i]:.3f}",
                "PASS" if ok_sv else "FAIL",
                f"{mui_:.2f}", f"{cap:.2f}", f"{dcr_m:.4f}",
                "PASS" if abs(mui_)<=cap else "FAIL",
                f"{vui_:.2f}", f"{pVi_:.2f}", f"{dcr_v:.4f}",
                "PASS" if vui_<=pVi_ else "FAIL",
            ])
        tbl(["x(m)",
             "σ_top Tr","σ_bot Tr","Transfer",
             "σ_top Sv","σ_bot Sv","Service",
             "Mu","φMn","DCR_M","Flexure",
             "Vu","φVn","DCR_V","Shear"],
            sum_rows,
            cw=[1.2,1.6,1.6,1.4,1.6,1.6,1.4,1.6,1.6,1.4,1.4,1.6,1.6,1.4,1.4])
        blank()

        h1("7.  Conclusion")
        all_pass = all(
            R["lim_tr_c"]<=R["tr_top"][i]<=R["lim_tr_t"] and
            R["lim_tr_c"]<=R["tr_bot"][i]<=R["lim_tr_t"] and
            R["sv1_top"][i] >= R["lim_sv_ct"] and
            R["sv1_bot"][i] >= R["lim_sv_ct"] and
            R["sv1_top"][i] <= R["lim_sv_t"]  and
            R["sv1_bot"][i] <= R["lim_sv_t"]  and
            abs(float(R["mu"][i])) <= max(float(R["phi_Mn_pos"][i]),
                                          abs(float(R["phi_Mn_neg"][i]))) and
            float(R["vu"][i]) <= float(R["phi_Vn"][i])
            for i in sta_idx
        )
        if all_pass:
            para("► OVERALL: The top flange tendon design is ADEQUATE for all "
                 "AASHTO LRFD limit states checked.",
                 bold=True, color=C_GREEN)
        else:
            para("► OVERALL: The design does NOT satisfy all limit states. "
                 "Revise tendon layout, spacing, or section geometry.",
                 bold=True, color=C_RED)
        blank()
        para("─── END OF CALCULATION ───", color=C_GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ─────────────────────────────────────────────────────────────────
    # EXPORT REPORT BUTTON — PROMINENT in main area + sidebar
    # ─────────────────────────────────────────────────────────────────
    try:
        report_bytes = make_report()

        # ── Main area prominent export button ─────────────────────────
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #0f2040, #1a3355, #2d5a8e);
            border-radius: 12px;
            padding: 20px 24px;
            margin-bottom: 24px;
            border: 2px solid #2d5a8e;
            box-shadow: 0 4px 20px rgba(217,119,6,0.35);
            display: flex;
            align-items: center;
            gap: 16px;
        ">
            <div style="font-size:2rem;">📥</div>
            <div>
                <div style="font-size:1.0rem; font-weight:800; color:#fde68a; line-height:1.2;">
                    Export Calculation Report (.docx)
                </div>
                <div style="font-size:0.78rem; color:#fcd34d; margin-top:3px; font-weight:500;">
                    รายงานการคำนวณฉบับสมบูรณ์ พร้อม Step-by-step derivation · ทุก Station
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        col_exp, col_pad = st.columns([1, 2])
        with col_exp:
            st.download_button(
                label="📥  Export Calculation Report  (.docx)",
                data=report_bytes,
                file_name=f"CalcReport_{proj_name.replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
            )

        # ── Sidebar export button ──────────────────────────────────────
        with st.sidebar:
            st.markdown("---")
            st.markdown("""
            <div style="text-align:center; color:#fcd34d; font-size:0.72rem;
                        font-weight:700; letter-spacing:0.1em; text-transform:uppercase;
                        margin-bottom:6px;">
                📥  Export Calculation Report
            </div>
            """, unsafe_allow_html=True)
            # Re-generate for sidebar button (separate BytesIO)
            report_bytes2 = make_report()
            st.download_button(
                label="📥  Export Report  (.docx)",
                data=report_bytes2,
                file_name=f"CalcReport_{proj_name.replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="sidebar_export",
            )

    except Exception as rep_err:
        st.error(f"Report generation error: {rep_err}")

    st.markdown("---")

    # ─────────────────────────────────────────────────────────────────
    # 6.  RESULT TABS  — UNCHANGED logic, improved styling
    # ─────────────────────────────────────────────────────────────────
    def dcr_style(df_in, col):
        def _s(val):
            try: v = float(val)
            except: return ""
            if v <= 0.80: return "background-color:#c6efce;color:#276221"
            if v <= 1.00: return "background-color:#ffeb9c;color:#9c6500"
            return "background-color:#ffc7ce;color:#9c0006"
        return df_in.style.map(_s, subset=[col])

    tabs = st.tabs([
        "📐 Geometry",
        "📉 Prestress Losses",
        "🚀 Transfer Stress",
        "⚖️ Service Stress",
        "💪 Flexure (Envelope)",
        "🔪 Shear",
        "📋 Summary",
    ])

    # ── TAB 0: Geometry ───────────────────────────────────────────────
    with tabs[0]:
        st.subheader("Top Flange Cross-Section with Tendon Layout")

        x_m    = R["x"]
        top_mm = np.zeros(N)
        bot_mm = -R["t"] * 1000.0
        cg_mm  = -R["yc"] * 1000.0
        tdn_mm = -R["z"] * 1000.0

        t_max_mm = float(R["t"].max()) * 1000.0
        t_min_mm = float(R["t"].min()) * 1000.0

        scale_k  = max(1.0, round(0.15 * width / (t_max_mm / 1000.0)))
        y_margin = t_max_mm * 1.8
        y_range  = [-t_max_mm - y_margin, y_margin]

        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=np.concatenate([x_m, x_m[::-1]]),
            y=np.concatenate([top_mm, bot_mm[::-1]]),
            fill="toself",
            fillcolor="rgba(173, 204, 240, 0.45)",
            line=dict(color="steelblue", width=1.5),
            name="Top Flange", hoverinfo="skip",
        ))
        fig.add_trace(go.Scatter(
            x=x_m, y=cg_mm, mode="lines",
            line=dict(color="gray", dash="dot", width=1),
            name="Section CG",
        ))
        fig.add_trace(go.Scatter(
            x=x_m, y=tdn_mm, mode="lines",
            line=dict(color="red", width=2.0),
            name="Tendon CGS", showlegend=True,
        ))
        _tdn_prep = prep(df_tdn)
        tdn_dot_x = _tdn_prep["x (m)"].values
        tdn_dot_y = -_tdn_prep["z_top (m)"].values * 1000.0
        fig.add_trace(go.Scatter(
            x=tdn_dot_x, y=tdn_dot_y, mode="markers",
            marker=dict(color="red", size=9, symbol="circle",
                        line=dict(color="white", width=1.5)),
            name="Tendon input pts", showlegend=True,
        ))
        for x_edge, label, a_pos in [
            (0.0,   "Edge L.Flange", "top right"),
            (width, "Edge R.Flange", "top left"),
        ]:
            fig.add_vline(
                x=x_edge,
                line=dict(color="rgba(0,170,170,0.85)", dash="dot", width=1.8),
                annotation_text=f"<b>{label}</b>",
                annotation_position=a_pos,
                annotation_font=dict(size=10, color="rgba(0,150,150,1)"),
            )
        for x_wf, label, a_pos in [
            (cl_lweb, "CL. L.Web", "top right"),
            (cl_rweb, "CL. R.Web", "top left"),
        ]:
            fig.add_vline(
                x=x_wf,
                line=dict(color="rgba(200,100,0,0.9)", dash="dash", width=2.0),
                annotation_text=f"<b>{label}</b>",
                annotation_position=a_pos,
                annotation_font=dict(size=10, color="rgba(200,100,0,1)"),
            )
        fig.update_layout(
            title="Top Flange Cross-Section with Tendon Layout",
            height=420,
            xaxis=dict(
                title="Distance from Left Edge (m)",
                range=[-width*0.04, width*1.04],
                showgrid=True, gridcolor="rgba(200,200,200,0.4)",
            ),
            yaxis=dict(
                title="Depth (mm)",
                range=y_range,
                showgrid=True, gridcolor="rgba(200,200,200,0.4)"
            ),
            legend=dict(orientation="h", y=-0.18),
            plot_bgcolor="white",
            margin=dict(t=50, b=80),
        )
        st.plotly_chart(fig, use_container_width=True)

        col_inf1, col_inf2, col_inf3, col_inf4 = st.columns(4)
        col_inf1.info(f"Scale y:x = 1:{int(scale_k)}")
        col_inf2.info(f"t_min = {t_min_mm:.0f} mm")
        col_inf3.info(f"CL.L.Web = {cl_lweb:.2f} m")
        col_inf4.info(f"CL.R.Web = {cl_rweb:.2f} m")

        c1, c2, c3 = st.columns(3)
        c1.metric("Aps (1m strip)", f"{R['Aps']*1e6:.2f} mm²")
        c2.metric("Pi", f"{R['Pi']:.2f} kN/m")
        c3.metric("Pe", f"{R['Pe']:.2f} kN/m")

    # ── TAB 1: Prestress Losses ───────────────────────────────────────
    with tabs[1]:
        st.subheader("📉 Prestress Loss Summary  (AASHTO LRFD 5.9.3)")
        _L = R["L"]
        _fpj = _L["fpj"]

        col_a, col_b, col_c = st.columns(3)
        col_a.metric("Jacking Stress fpj",  f"{_fpj:.1f} MPa")
        col_b.metric("Aps (1m strip)",       f"{_L['Aps']*1e6:.2f} mm²")
        col_c.metric("Pj (per 1m strip)",    f"{_L['Pj']:.2f} kN/m")

        st.markdown("---")
        st.markdown("#### Immediate Losses")
        _d = {
            "Loss Component": [
                "1. Friction  ΔfpF  (at midspan)",
                "2. Anchorage Set  ΔfpA  (at jacking end)",
                "3. Elastic Shortening  ΔfpES",
                "Total Immediate  Δfi",
            ],
            "Formula  [AASHTO Ref.]": [
                "fpj(1−e^(−μα−Kx))  [5.9.3.2.1]",
                "Δ·Ep/Lpa  [5.9.3.2.2]",
                "(Ep/Eci)·fcgp  [5.9.3.2.3]",
                "ΔfpF + ΔfpA + ΔfpES",
            ],
            "Key Params": [
                f"μ={0.20}, K={0.0066}, α={_L['alpha']:.4f}rad",
                f"Δ={st.session_state.get('anch_slip_mm',6):.0f}mm, Lpa={_L['Lpa']:.2f}m",
                f"fcgp={_L['fcgp']:.3f}MPa, Eci={_L['Eci']:.0f}MPa",
                "",
            ],
            "Loss (MPa)": [
                f"{_L['delta_fpF']:.2f}",
                f"{_L['delta_fpA']:.2f}",
                f"{_L['delta_fpES']:.2f}",
                f"{_L['delta_imm']:.2f}",
            ],
            "% of fpj": [
                f"{_L['delta_fpF']/_fpj*100:.2f}",
                f"{_L['delta_fpA']/_fpj*100:.2f}",
                f"{_L['delta_fpES']/_fpj*100:.2f}",
                f"{_L['imm_loss_pct']:.2f}",
            ],
        }
        st.dataframe(pd.DataFrame(_d), use_container_width=True)

        col_x, col_y = st.columns(2)
        col_x.metric("fpi (after imm. losses)", f"{_L['fpi_eff']:.1f} MPa")
        col_y.metric("Pi (per 1m strip)",        f"{_L['Pi']:.2f} kN/m")

        st.markdown("---")
        st.markdown("#### Long-Term Losses  (Approximate Method)")
        _d2 = {
            "Loss Component": [
                "4. Shrinkage  ΔfpSH",
                "5. Creep  ΔfpCR",
                "6. Relaxation  ΔfpR2",
                "Total Long-term  ΔfLT",
                "TOTAL LOSS  Δftotal",
            ],
            "Formula  [AASHTO Ref.]": [
                "εbdf × Ep  [5.9.3.4.2a]",
                "(Ep/Ec) × fcgp × ψb  [5.9.3.4.2b]",
                "(fpt/KL)(fpt/fpy − 0.55)  [5.9.3.4.3c]",
                "ΔfpSH + ΔfpCR + ΔfpR",
                "Δfi + ΔfLT",
            ],
            "Key Params": [
                f"εbdf={_L['eps_bdf']:.5f}, kvs={_L['kvs']:.3f}, khs={_L['khs']:.3f}, kf={_L['kf']:.3f}",
                f"ψb={_L['psi_b']:.3f}, fcgp={_L['fcgp_lt']:.3f}MPa, khc={_L['khc']:.3f}",
                f"KL={45}, fpt/fpy={_L['fpi_eff']/(_fpj*0.9):.3f}",
                "", "",
            ],
            "Loss (MPa)": [
                f"{_L['delta_fpSH']:.2f}",
                f"{_L['delta_fpCR']:.2f}",
                f"{_L['delta_fpR']:.2f}",
                f"{_L['delta_lt']:.2f}",
                f"{_L['delta_imm']+_L['delta_lt']:.2f}",
            ],
            "% of fpj": [
                f"{_L['delta_fpSH']/_fpj*100:.2f}",
                f"{_L['delta_fpCR']/_fpj*100:.2f}",
                f"{_L['delta_fpR']/_fpj*100:.2f}",
                f"{_L['lt_loss_pct']:.2f}",
                f"{_L['total_loss_pct']:.2f}",
            ],
        }
        st.dataframe(pd.DataFrame(_d2), use_container_width=True)

        col_p, col_q, col_r = st.columns(3)
        col_p.metric("fpe (effective)", f"{_L['fpe']:.1f} MPa",
                     delta=f"{_L['fpe']/_fpj*100:.1f}% of fpj")
        col_q.metric("Pe (per 1m strip)", f"{_L['Pe']:.2f} kN/m")
        col_r.metric("Total Loss", f"{_L['total_loss_pct']:.2f}%",
                     delta=f"{_L['fpe']:.0f} MPa remaining")

        st.markdown("---")
        st.markdown("#### Key Factors Used")
        _factors = {
            "Parameter (description)": [
                "Ec  (modulus of elasticity of concrete at service)",
                "Eci  (modulus of elasticity of concrete at transfer)",
                "Ep  (modulus of elasticity of prestress steel)",
                "μ  (friction coefficient, grouted duct)",
                "K  (wobble coefficient, grouted duct)",
                "α  (total angular change of tendon)",
                "L  (tendon length along profile)",
                "friction slope  (rate of friction loss per unit length)",
                "Lpa  (anchor set influence length)",
                "Δ  (anchorage slip)",
                "V/S  (volume-to-surface ratio)",
                "kvs  (V/S size factor for shrinkage/creep)",
                "khs  (humidity factor for shrinkage)",
                "khc  (humidity factor for creep)",
                "kf  (concrete strength factor, uses f'ci in ksi)",
                "ψb  (creep coefficient)",
                "εbdf  (shrinkage strain)",
            ],
            "Value  [Unit]": [
                f"{_L['Ec']:.0f} MPa",
                f"{_L['Eci']:.0f} MPa",
                "197,000 MPa",
                "0.20  [-]",
                "0.0066 rad/m",
                f"{_L['alpha']:.4f} rad",
                f"{_L['L_ten']:.3f} m",
                f"{_L['friction_slope']:.4f} MPa/m",
                f"{_L['Lpa']:.3f} m",
                f"{st.session_state.get('anch_slip_mm', 6.0):.1f} mm",
                f"{_L['VS_mm']:.1f} mm",
                f"{_L['kvs']:.4f}",
                f"{_L['khs']:.4f}",
                f"{_L['khc']:.4f}",
                f"{_L['kf']:.4f}  (f'ci_ksi = {_L['fci_ksi']:.3f} ksi)",
                f"{_L['psi_b']:.4f}",
                f"{_L['eps_bdf']:.6f}",
            ],
        }
        st.dataframe(pd.DataFrame(_factors), use_container_width=True)

    # ── TAB 2: Transfer Stress ────────────────────────────────────────
    with tabs[2]:
        st.subheader("Stress Check — Transfer  (Pi + M_DL  |  Net section)")
        fig2 = go.Figure([
            go.Scatter(x=R["x"], y=R["tr_top"], name="Top",    line_color="red"),
            go.Scatter(x=R["x"], y=R["tr_bot"], name="Bottom", line_color="blue"),
        ])
        fig2.add_hline(y=R["lim_tr_c"], line_dash="dash", line_color="orange",
                       annotation_text=f"−0.60f'ci = {R['lim_tr_c']:.2f} MPa")
        fig2.add_hline(y=R["lim_tr_t"], line_dash="dash", line_color="green",
                       annotation_text=f"+0.62√f'ci = +{R['lim_tr_t']:.3f} MPa")
        fig2.update_layout(height=380, xaxis_title="x (m)", yaxis_title="Stress (MPa)")
        st.plotly_chart(fig2, use_container_width=True)
        rows_tr = [{"x (m)": f"{R['x'][i]:.2f}",
                    "σ_top (MPa)": f"{R['tr_top'][i]:.4f}",
                    "σ_bot (MPa)": f"{R['tr_bot'][i]:.4f}",
                    "Status": "✅" if (R["lim_tr_c"]<=R["tr_top"][i]<=R["lim_tr_t"] and
                                       R["lim_tr_c"]<=R["tr_bot"][i]<=R["lim_tr_t"]) else "❌"}
                   for i in sta_idx]
        st.dataframe(pd.DataFrame(rows_tr), use_container_width=True)

    # ── TAB 3: Service Stress ─────────────────────────────────────────
    with tabs[3]:
        st.subheader("Stress Check — Service I  (Pe + Ms1  |  Gross section)")
        fig3 = go.Figure([
            go.Scatter(x=R["x"], y=R["sv1_top"], name="Top",    line_color="red"),
            go.Scatter(x=R["x"], y=R["sv1_bot"], name="Bottom", line_color="blue"),
        ])
        fig3.add_hline(y=R["lim_sv_ct"], line_dash="dash", line_color="orange",
                       annotation_text=f"−0.60f'c = {R['lim_sv_ct']:.2f} MPa")
        fig3.add_hline(y=R["lim_sv_cp"], line_dash="dot", line_color="goldenrod",
                       annotation_text=f"−0.45f'c = {R['lim_sv_cp']:.2f} MPa")
        fig3.add_hline(y=R["lim_sv_t"],  line_dash="dash", line_color="green",
                       annotation_text=f"+0.50√f'c = +{R['lim_sv_t']:.3f} MPa")
        fig3.update_layout(height=380, xaxis_title="x (m)", yaxis_title="Stress (MPa)")
        st.plotly_chart(fig3, use_container_width=True)
        rows_sv = [{"x (m)":       f"{R['x'][i]:.2f}",
                    "σ_top (MPa)": f"{R['sv1_top'][i]:.4f}",
                    "σ_bot (MPa)": f"{R['sv1_bot'][i]:.4f}",
                    "Comp. Limit": f"{R['lim_sv_ct']:.2f}",
                    "Tens. Limit": f"+{R['lim_sv_t']:.3f}",
                    "Status": "✅" if (R["sv1_top"][i] >= R["lim_sv_ct"] and
                                       R["sv1_bot"][i] >= R["lim_sv_ct"] and
                                       R["sv1_top"][i] <= R["lim_sv_t"]  and
                                       R["sv1_bot"][i] <= R["lim_sv_t"]) else "❌"}
                   for i in sta_idx]
        st.dataframe(pd.DataFrame(rows_sv), use_container_width=True)

    # ── TAB 4: Flexure ────────────────────────────────────────────────
    with tabs[4]:
        st.subheader("Flexural Strength Envelope  —  Strength I")
        fig4 = go.Figure()
        fig4.add_trace(go.Scatter(x=R["x"], y=R["phi_Mn_pos"], name="+φMn",
                                   line=dict(color="green", dash="dash", width=2)))
        fig4.add_trace(go.Scatter(x=R["x"], y=R["phi_Mn_neg"], name="−φMn",
                                   line=dict(color="darkgreen", dash="dash", width=2)))
        fig4.add_trace(go.Scatter(x=R["x"], y=R["mu"], name="Mu",
                                   fill="tozeroy", line_color="rgba(220,50,50,0.8)"))
        fig4.update_layout(height=380, xaxis_title="x (m)", yaxis_title="Moment (kNm/m)")
        st.plotly_chart(fig4, use_container_width=True)
        rows_flx = []
        for i in sta_idx:
            mx   = float(R["mu"][i])
            cap  = float(R["phi_Mn_pos"][i]) if mx>=0 else abs(float(R["phi_Mn_neg"][i]))
            cdp  = float(R["cdp_pos"][i])    if mx>=0 else float(R["cdp_neg"][i])
            dcr  = abs(mx)/cap if cap>0 else 999
            rows_flx.append({"x (m)": f"{R['x'][i]:.2f}",
                              "Mu (kNm/m)":  f"{mx:.4f}",
                              "φMn (kNm/m)": f"{cap:.4f}",
                              "DCR":         f"{dcr:.4f}",
                              "c/dp":        f"{cdp:.4f}",
                              "Strength":    "✅" if abs(mx)<=cap else "❌",
                              "Ductility":   "✅" if cdp<=0.42   else "❌"})
        df_flx = pd.DataFrame(rows_flx)
        st.dataframe(dcr_style(df_flx, "DCR"), use_container_width=True)

    # ── TAB 5: Shear ──────────────────────────────────────────────────
    with tabs[5]:
        st.subheader("Shear Strength  —  Strength I  (β=2.0)")
        fig5 = go.Figure([
            go.Scatter(x=R["x"], y=R["phi_Vn"], name="φVn",
                       line=dict(color="green", width=2)),
            go.Scatter(x=R["x"], y=R["vu"],     name="Vu  (factored)",
                       fill="tozeroy", line_color="rgba(0,100,220,0.8)"),
        ])
        fig5.update_layout(height=380, xaxis_title="x (m)", yaxis_title="Shear (kN/m)")
        st.plotly_chart(fig5, use_container_width=True)
        rows_shr = []
        for i in sta_idx:
            vui_= float(R["vu"][i]);  pVi_= float(R["phi_Vn"][i])
            dcr = vui_/pVi_ if pVi_>0 else 999
            rows_shr.append({"x (m)":      f"{R['x'][i]:.2f}",
                              "dv (mm)":    f"{R['dv'][i]*1000:.2f}",
                              "Vc (kN/m)":  f"{R['Vc'][i]:.4f}",
                              "φVn (kN/m)": f"{pVi_:.4f}",
                              "Vu (kN/m)":  f"{vui_:.4f}",
                              "DCR":        f"{dcr:.4f}",
                              "Status":     "✅" if vui_<=pVi_ else "❌"})
        df_shr = pd.DataFrame(rows_shr)
        st.dataframe(dcr_style(df_shr, "DCR"), use_container_width=True)

    # ── TAB 6: Summary ────────────────────────────────────────────────
    with tabs[6]:
        st.subheader("📋 Overall Design Summary")
        rows_sum = []
        for i in sta_idx:
            mui_= float(R["mu"][i]);  vui_= float(R["vu"][i])
            cap = (float(R["phi_Mn_pos"][i]) if mui_>=0
                   else abs(float(R["phi_Mn_neg"][i])))
            pVi_= float(R["phi_Vn"][i])
            ok_tr = (R["lim_tr_c"]<=R["tr_top"][i]<=R["lim_tr_t"] and
                     R["lim_tr_c"]<=R["tr_bot"][i]<=R["lim_tr_t"])
            ok_sv = (R["sv1_top"][i] >= R["lim_sv_ct"] and
                     R["sv1_bot"][i] >= R["lim_sv_ct"] and
                     R["sv1_top"][i] <= R["lim_sv_t"]  and
                     R["sv1_bot"][i] <= R["lim_sv_t"])
            dcr_m = abs(mui_)/cap   if cap >0 else 999
            dcr_v = vui_/pVi_       if pVi_>0 else 999
            rows_sum.append({
                "x (m)":     f"{R['x'][i]:.2f}",
                "Transfer":  "✅" if ok_tr else "❌",
                "Service":   "✅" if ok_sv else "❌",
                "DCR_M":     f"{dcr_m:.4f}",
                "Flexure":   "✅" if abs(mui_)<=cap  else "❌",
                "DCR_V":     f"{dcr_v:.4f}",
                "Shear":     "✅" if vui_<=pVi_       else "❌",
            })
        df_sum = pd.DataFrame(rows_sum)
        st.dataframe(dcr_style(df_sum, "DCR_M"), use_container_width=True)

        all_ok = all(
            r["Transfer"]=="✅" and r["Service"]=="✅" and
            r["Flexure"]=="✅"  and r["Shear"]=="✅"
            for r in rows_sum
        )
        if all_ok:
            st.success("✅  DESIGN ADEQUATE — All checks pass at all stations.")
        else:
            st.error("❌  DESIGN INADEQUATE — One or more checks fail. Revise design.")

        st.caption("DCR: 🟢 ≤0.80  |  🟡 0.80–1.00  |  🔴 >1.00")

except Exception as err:
    st.error(f"Calculation error: {err}")
    raise